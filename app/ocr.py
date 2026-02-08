import io
import os
import sys
# Add the current directory to Python path to ensure imports work
current_dir = os.path.dirname(os.path.abspath(__file__))
if current_dir not in sys.path:
    sys.path.insert(0, current_dir)

import shutil
import tempfile
import uuid
import re
import time
import hashlib
import zipfile
from datetime import datetime, timedelta
from concurrent.futures import ThreadPoolExecutor
from functools import lru_cache

from flask import Flask, render_template, send_file, request, abort, Response, jsonify, send_from_directory, after_this_request
from werkzeug.utils import secure_filename

from PyPDF2 import PdfReader, PdfWriter, PdfMerger
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4, letter
from PIL import Image
from pdfminer.high_level import extract_text

# -----------------------------------------------------------------------------
# OCR IMPORTS - COMPLETE SET
# -----------------------------------------------------------------------------
try:
    from ocr import (
        pdf_to_word_with_ocr,
        pdf_to_text_with_ocr,
        image_to_text,
        image_to_word,
        extract_text_from_file,
        is_scanned_pdf,
        is_image_based_document,
        ocr_pdf_to_searchable_pdf
    )
    OCR_AVAILABLE = True
    print("✅ OCR module loaded successfully")
except ImportError as e:
    print(f"⚠️ OCR module not available: {e}")
    OCR_AVAILABLE = False
    # Define dummy functions
    def pdf_to_word_with_ocr(*args, **kwargs):
        raise ImportError("OCR not available")
    def pdf_to_text_with_ocr(*args, **kwargs):
        raise ImportError("OCR not available")
    def image_to_text(*args, **kwargs):
        raise ImportError("OCR not available")
    def image_to_word(*args, **kwargs):
        raise ImportError("OCR not available")
    def extract_text_from_file(*args, **kwargs):
        raise ImportError("OCR not available")
    def is_scanned_pdf(*args, **kwargs):
        return False
    def is_image_based_document(*args, **kwargs):
        return False
    def ocr_pdf_to_searchable_pdf(*args, **kwargs):
        raise ImportError("OCR not available")

# -----------------------------------------------------------------------------
# Flask app configuration with correct template paths
# -----------------------------------------------------------------------------
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

app = Flask(__name__, 
           template_folder=os.path.join(BASE_DIR, 'templates'),
           static_folder=os.path.join(BASE_DIR, 'static'))

print(f"✅ App initialized with:")
print(f"   Base directory: {BASE_DIR}")
print(f"   Templates: {os.path.join(BASE_DIR, 'templates')}")
print(f"   Static files: {os.path.join(BASE_DIR, 'static')}")
print(f"   OCR Available: {OCR_AVAILABLE}")

# -----------------------------------------------------------------------------
# Performance settings
# -----------------------------------------------------------------------------
MAX_CONTENT_LENGTH = 50 * 1024 * 1024  # 50 MB per file
UPLOAD_DIR = os.path.join(tempfile.gettempdir(), "imasterpdf_uploads")
OUTPUT_DIR = os.path.join(tempfile.gettempdir(), "imasterpdf_outputs")
CLEANUP_AGE_MINUTES = 30
MAX_WORKERS = 4
MAX_PAGES_TO_EXTRACT = 100
CACHE_ENABLED = True
OCR_ENABLED = True  # Enable/disable OCR globally

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH

executor = ThreadPoolExecutor(max_workers=MAX_WORKERS)
conversion_cache = {}

ALLOWED_IMAGE_EXT = {'.jpg', '.jpeg', '.png', '.webp', '.bmp', '.tiff', '.tif'}
ALLOWED_PDF_EXT = {'.pdf'} 
ALLOWED_WORD_EXT = {'.docx', '.doc'}
ALLOWED_TEXT_EXT = {'.txt'}

# -----------------------------------------------------------------------------
# Helper Functions
# -----------------------------------------------------------------------------
def clean_text_for_xml(text):
    if not text:
        return ""
    
    text = text.replace('\x00', '')
    text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]', '', text)
    
    replacements = {
        '\u2028': ' ',
        '\u2029': ' ',
        '\uFEFF': '',
    }
    
    for old, new in replacements.items():
        text = text.replace(old, new)
    
    try:
        text = text.encode('utf-8', 'ignore').decode('utf-8')
    except:
        text = text.encode('ascii', 'ignore').decode('ascii')
    
    return text

def safe_add_paragraph(doc, text):
    try:
        cleaned_text = clean_text_for_xml(text)
        if cleaned_text.strip():
            doc.add_paragraph(cleaned_text.strip())
    except Exception:
        pass

def get_file_hash(file_path):
    hasher = hashlib.md5()
    with open(file_path, 'rb') as f:
        buf = f.read(65536)
        while len(buf) > 0:
            hasher.update(buf)
            buf = f.read(65536)
    return hasher.hexdigest()

def fast_extract_text(pdf_path):
    start_time = time.time()
    file_size = os.path.getsize(pdf_path)
    
    if file_size < 102400:
        try:
            with open(pdf_path, 'rb') as file:
                reader = PdfReader(file)
                text = ""
                for page in reader.pages[:MAX_PAGES_TO_EXTRACT]:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                if text.strip():
                    print(f"Fast PyPDF2 extraction: {time.time() - start_time:.2f}s")
                    return clean_text_for_xml(text)
        except:
            pass
    
    try:
        def extract_page(page):
            try:
                return page.extract_text() or ""
            except:
                return ""
        
        with open(pdf_path, 'rb') as file:
            reader = PdfReader(file)
            pages = reader.pages[:MAX_PAGES_TO_EXTRACT]
            texts = list(executor.map(extract_page, pages))
            result = "\n".join(texts)
            if result and len(result.strip()) > 50:
                print(f"Parallel extraction: {time.time() - start_time:.2f}s")
                return clean_text_for_xml(result)
    except:
        pass
    
    try:
        text = extract_text(
            pdf_path,
            maxpages=MAX_PAGES_TO_EXTRACT,
            caching=True,
            laparams=None
        )
        print(f"Optimized pdfminer extraction: {time.time() - start_time:.2f}s")
        return clean_text_for_xml(text or "")
    except Exception as e:
        print(f"All extraction methods failed: {e}")
        return ""

def optimize_pdf_for_extraction(pdf_path):
    try:
        file_size = os.path.getsize(pdf_path)
        if file_size < 5 * 1024 * 1024:
            return pdf_path
            
        with open(pdf_path, 'rb') as file:
            reader = PdfReader(file)
            if len(reader.pages) <= MAX_PAGES_TO_EXTRACT:
                return pdf_path
            
            writer = PdfWriter()
            for i in range(min(MAX_PAGES_TO_EXTRACT, len(reader.pages))):
                writer.add_page(reader.pages[i])
            
            optimized_path = pdf_path + "_optimized.pdf"
            with open(optimized_path, 'wb') as f:
                writer.write(f)
            
            return optimized_path
    except:
        return pdf_path

def ext_of(filename):
    return os.path.splitext(filename.lower())[1]

def validate_file(stream):
    stream.seek(0, os.SEEK_END)
    size = stream.tell()
    stream.seek(0)
    if size < 1024:
        abort(Response("File too small (min 1 KB).", status=400))
    if size > MAX_CONTENT_LENGTH:
        abort(Response("File too large (max 50 MB).", status=400))

def generate_unique_filename(original_filename, suffix=""):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    unique_id = str(uuid.uuid4())[:12]
    name, ext = os.path.splitext(original_filename)
    safe_name = secure_filename(name)
    
    if suffix:
        return f"{safe_name}_{suffix}_{timestamp}_{unique_id}{ext}"
    return f"{safe_name}_{timestamp}_{unique_id}{ext}"

def save_uploads(files):
    saved = []
    for storage in files:
        validate_file(storage.stream)
        filename = secure_filename(storage.filename)
        if not filename:
            abort(Response("Invalid filename.", status=400))
        
        unique_filename = generate_unique_filename(storage.filename)
        path = os.path.join(UPLOAD_DIR, unique_filename)
        storage.save(path)
        saved.append(path)
    return saved

def cleanup_temp():
    cutoff = datetime.utcnow() - timedelta(minutes=CLEANUP_AGE_MINUTES)
    for base in (UPLOAD_DIR, OUTPUT_DIR):
        for name in os.listdir(base):
            path = os.path.join(base, name)
            try:
                mtime = datetime.utcfromtimestamp(os.path.getmtime(path))
                if mtime < cutoff:
                    if os.path.isdir(path):
                        shutil.rmtree(path, ignore_errors=True)
                    else:
                        os.remove(path)
            except Exception:
                pass

def wrap_text(text, max_chars=95):
    words = text.split(' ')
    lines, current = [], []
    length = 0
    for w in words:
        add_len = len(w) + (1 if current else 0)
        if length + add_len <= max_chars:
            current.append(w)
            length += add_len
        else:
            lines.append(' '.join(current))
            current = [w]
            length = len(w)
    if current:
        lines.append(' '.join(current))
    return lines

def parse_pages(pages_str):
    pages = set()
    parts = [p.strip() for p in pages_str.split(',') if p.strip()]
    for part in parts:
        if '-' in part:
            a, b = part.split('-', 1)
            try:
                start = int(a); end = int(b)
                for i in range(min(start, end), max(start, end)+1):
                    pages.add(i)
            except ValueError:
                abort(Response("Invalid page range.", status=400))
        else:
            try:
                pages.add(int(part))
            except ValueError:
                abort(Response("Invalid page number.", status=400))
    return pages

def safe_remove(path):
    try:
        if os.path.exists(path):
            os.remove(path)
    except Exception:
        pass

def safe_remove_all(paths):
    for path in paths:
        safe_remove(path)

# -----------------------------------------------------------------------------
# New OCR Helper Functions
# -----------------------------------------------------------------------------
def should_use_ocr(file_path, file_type):
    """
    Determine whether to use OCR for a file
    """
    if not OCR_ENABLED or not OCR_AVAILABLE:
        return False
    
    if file_type == 'pdf':
        return is_scanned_pdf(file_path)
    elif file_type == 'image':
        return True
    elif file_type == 'word':
        return is_image_based_document(file_path)
    
    return False

def extract_text_with_fallback(file_path, use_ocr=False):
    """
    Extract text with OCR fallback
    """
    if use_ocr and OCR_AVAILABLE:
        print("🔄 Using OCR for text extraction")
        try:
            return extract_text_from_file(file_path)
        except Exception as e:
            print(f"⚠️ OCR extraction failed: {e}")
    
    # Try regular extraction
    try:
        if file_path.lower().endswith('.pdf'):
            return fast_extract_text(file_path)
        elif file_path.lower().endswith(('.docx', '.doc')):
            doc = Document(file_path)
            text = ""
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            return text
    except Exception as e:
        print(f"⚠️ Regular extraction failed: {e}")
    
    return ""

# -----------------------------------------------------------------------------
# Route Definitions
# -----------------------------------------------------------------------------
@app.route('/')
@app.route('/index')
@app.route('/index.html')
def index():
    return render_template('index.html')

@app.route('/split')
@app.route('/split.html')
def split_pdf():
    return render_template('split.html')

@app.route('/mergepdf')
@app.route('/mergepdf.html')
def merge_pdf():
    return render_template('mergepdf.html')

@app.route('/deletepdf')
@app.route('/deletepdf.html')
def delete_pdf():
    return render_template('deletepdf.html')

@app.route('/rotatepdf')
@app.route('/rotatepdf.html')
def rotate_pdf():
    return render_template('rotatepdf.html')

@app.route('/pdftoword')
@app.route('/pdftoword.html')
def pdf_to_word_page():
    return render_template('pdftoword.html')

@app.route('/lockpdf')
@app.route('/lockpdf.html')
def lock_pdf():
    return render_template('lockpdf.html')

@app.route('/unlockpdf')
@app.route('/unlockpdf.html')
def unlock_pdf():
    return render_template('unlockpdf.html')

@app.route('/wordtopdf')
@app.route('/wordtopdf.html')
def word_to_pdf_page():
    return render_template('wordtopdf.html')

@app.route('/mergeword')
@app.route('/mergeword.html')
def merge_word():
    return render_template('mergeword.html')

@app.route('/wordtotext')
@app.route('/wordtotext.html')
def word_to_text():
    return render_template('wordtotext.html')

@app.route('/texttopdf')
@app.route('/texttopdf.html')
def text_to_pdf():
    return render_template('texttopdf.html')

@app.route('/texttoword')
@app.route('/texttoword.html')
def text_to_word():
    return render_template('texttoword.html')

@app.route('/imagestopdf')
@app.route('/imagestopdf.html')
def images_to_pdf():
    return render_template('imagestopdf.html')

@app.route('/<path:filename>.html')
def serve_html(filename):
    try:
        return render_template(f'{filename}.html')
    except:
        abort(404)

# -----------------------------------------------------------------------------
# API Routes
# -----------------------------------------------------------------------------
@app.route('/api/contact', methods=['POST'])
def api_contact():
    data = request.get_json(silent=True) or {}
    name = (data.get('name') or '').strip()
    email = (data.get('email') or '').strip()
    message = (data.get('message') or '').strip()
    if not name or not email or not message:
        return jsonify({"error": "Please provide name, email, and message."}), 400
    return jsonify({"status": "ok", "received": {"name": name, "email": email}}), 200

# -----------------------------------------------------------------------------
# Enhanced PDF to Word with OCR Support
# -----------------------------------------------------------------------------
@app.route('/api/pdf-to-word', methods=['POST'])
def api_pdf_to_word():
    start_time = time.time()
    cleanup_temp()
    
    files = request.files.getlist('files')
    if not files or len(files) != 1:
        return jsonify({"error": "Upload exactly one PDF."}), 400
    
    paths = save_uploads(files)
    pdf_path = paths[0]
    
    if ext_of(pdf_path) not in ALLOWED_PDF_EXT:
        safe_remove(pdf_path)
        return jsonify({"error": "Only PDF files are allowed."}), 400

    try:
        # Check if OCR should be used
        use_ocr = should_use_ocr(pdf_path, 'pdf')
        print(f"📄 File: {os.path.basename(pdf_path)}, Use OCR: {use_ocr}")
        
        # Generate output name
        original_name = secure_filename(files[0].filename)
        output_name = generate_unique_filename(original_name, "converted")
        output_name = os.path.splitext(output_name)[0] + ".docx"
        
        # Method 1: Use OCR if needed
        if use_ocr and OCR_AVAILABLE:
            print("🔄 Converting PDF to Word using OCR...")
            try:
                # Use OCR-based conversion
                buffer = pdf_to_word_with_ocr(pdf_path)
                
                response = send_file(
                    buffer,
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    as_attachment=True,
                    download_name=output_name
                )
                
                @after_this_request
                def cleanup(response):
                    safe_remove(pdf_path)
                    return response
                
                print(f"✅ PDF to Word (OCR): {time.time() - start_time:.2f}s")
                return response
                
            except Exception as ocr_error:
                print(f"⚠️ OCR conversion failed: {ocr_error}")
                # Fall back to regular method
        
        # Method 2: Regular text extraction
        print("📝 Converting PDF to Word using text extraction...")
        optimized_path = optimize_pdf_for_extraction(pdf_path)
        text = fast_extract_text(optimized_path)
        
        if optimized_path != pdf_path:
            safe_remove(optimized_path)
        
        # Create Word document
        doc = Document()
        
        if text and text.strip():
            paragraphs = [p for p in text.split('\n\n') if p.strip()]
            if len(paragraphs) > 500:
                paragraphs = paragraphs[:500]
                doc.add_paragraph("[Document truncated - first 500 paragraphs shown]")
            
            for para in paragraphs:
                safe_add_paragraph(doc, para)
        else:
            # No text found, try OCR as last resort
            if OCR_AVAILABLE:
                print("⚠️ No text found, trying OCR...")
                try:
                    text = pdf_to_text_with_ocr(pdf_path, max_pages=50)
                    if text.strip():
                        for line in text.split('\n'):
                            if line.strip():
                                doc.add_paragraph(line.strip())
                    else:
                        doc.add_paragraph("No text could be extracted from this PDF.")
                except Exception as e:
                    doc.add_paragraph(f"Text extraction failed: {str(e)}")
            else:
                doc.add_paragraph("No text could be extracted from this PDF.")
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        response = send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=output_name
        )
        
        @after_this_request
        def cleanup(response):
            safe_remove(pdf_path)
            return response
        
        print(f"✅ PDF to Word: {time.time() - start_time:.2f}s")
        return response
        
    except Exception as e:
        safe_remove(pdf_path)
        return jsonify({"error": f"Conversion failed: {str(e)}"}), 500

# -----------------------------------------------------------------------------
# Enhanced Word to PDF with OCR Support
# -----------------------------------------------------------------------------
@app.route('/api/word-to-pdf', methods=['POST'])
def api_word_to_pdf():
    cleanup_temp()
    files = request.files.getlist('files')
    if not files or len(files) != 1:
        return jsonify({"error": "Upload exactly one Word file."}), 400

    paths = save_uploads(files)
    doc_path = paths[0]

    if ext_of(doc_path) not in ALLOWED_WORD_EXT:
        safe_remove(doc_path)
        return jsonify({"error": "Only DOC/DOCX files are supported."}), 400

    try:
        original_name = secure_filename(files[0].filename)
        output_name = generate_unique_filename(original_name, "converted")
        output_name = os.path.splitext(output_name)[0] + ".pdf"
        
        # Check if it's image-based
        use_ocr = should_use_ocr(doc_path, 'word')
        print(f"📄 Word file: {os.path.basename(doc_path)}, Use OCR: {use_ocr}")
        
        # Extract text with appropriate method
        text = ""
        if use_ocr and OCR_AVAILABLE:
            print("🔄 Extracting text from Word using OCR...")
            try:
                text = extract_text_from_file(doc_path)
            except Exception as e:
                print(f"⚠️ OCR extraction failed: {e}")
                # Fall back to regular extraction
                doc = Document(doc_path)
                paragraphs = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        cleaned = clean_text_for_xml(para.text)
                        if cleaned.strip():
                            paragraphs.append(cleaned)
                text = '\n\n'.join(paragraphs)
        else:
            # Regular extraction
            doc = Document(doc_path)
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    cleaned = clean_text_for_xml(para.text)
                    if cleaned.strip():
                        paragraphs.append(cleaned)
            text = '\n\n'.join(paragraphs)
        
        # If still no text, use OCR as fallback
        if not text or len(text.strip()) < 10:
            if OCR_AVAILABLE:
                print("⚠️ Little text found, trying OCR fallback...")
                try:
                    text = extract_text_from_file(doc_path) or "No text content"
                except:
                    text = "No text content"
            else:
                text = "No text content"
        
        # Create PDF
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        width, height = letter
        left_margin = 50
        top = height - 50
        line_height = 14
        
        if text:
            paragraphs = text.split('\n\n')
            for para in paragraphs[:200]:
                if para.strip():
                    lines = wrap_text(para, max_chars=95)
                    for line in lines:
                        c.drawString(left_margin, top, line)
                        top -= line_height
                        if top < 50:
                            c.showPage()
                            top = height - 50
                    top -= line_height / 2
                    if top < 50:
                        c.showPage()
                        top = height - 50
        
        c.save()
        buffer.seek(0)
        
        response = send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=output_name
        )
        
        @after_this_request
        def cleanup(response):
            safe_remove(doc_path)
            return response
        
        return response
        
    except Exception as e:
        safe_remove(doc_path)
        return jsonify({"error": f"Conversion failed: {str(e)}"}), 500

# -----------------------------------------------------------------------------
# NEW: Image to Text API with OCR
# -----------------------------------------------------------------------------
@app.route('/api/image-to-text', methods=['POST'])
def api_image_to_text():
    """Convert image to text using OCR"""
    if not OCR_AVAILABLE:
        return jsonify({"error": "OCR functionality is not available."}), 500
    
    cleanup_temp()
    files = request.files.getlist('files')
    if not files or len(files) != 1:
        return jsonify({"error": "Upload exactly one image."}), 400
    
    paths = save_uploads(files)
    image_path = paths[0]
    
    if ext_of(image_path) not in ALLOWED_IMAGE_EXT:
        safe_remove(image_path)
        return jsonify({"error": "Only image files are allowed."}), 400

    try:
        original_name = secure_filename(files[0].filename)
        output_name = generate_unique_filename(original_name, "extracted_text")
        output_name = os.path.splitext(output_name)[0] + ".txt"
        
        # Extract text using OCR
        text = image_to_text(image_path)
        
        if not text or not text.strip():
            text = "No text could be extracted from the image."
        
        buffer = io.BytesIO(text.encode('utf-8'))
        
        response = send_file(
            buffer,
            mimetype='text/plain',
            as_attachment=True,
            download_name=output_name
        )
        
        @after_this_request
        def cleanup(response):
            safe_remove(image_path)
            return response
        
        return response
        
    except Exception as e:
        safe_remove(image_path)
        return jsonify({"error": f"OCR failed: {str(e)}"}), 500

# -----------------------------------------------------------------------------
# NEW: Image to Word API with OCR
# -----------------------------------------------------------------------------
@app.route('/api/image-to-word', methods=['POST'])
def api_image_to_word():
    """Convert image to Word using OCR"""
    if not OCR_AVAILABLE:
        return jsonify({"error": "OCR functionality is not available."}), 500
    
    cleanup_temp()
    files = request.files.getlist('files')
    if not files or len(files) != 1:
        return jsonify({"error": "Upload exactly one image."}), 400
    
    paths = save_uploads(files)
    image_path = paths[0]
    
    if ext_of(image_path) not in ALLOWED_IMAGE_EXT:
        safe_remove(image_path)
        return jsonify({"error": "Only image files are allowed."}), 400

    try:
        original_name = secure_filename(files[0].filename)
        output_name = generate_unique_filename(original_name, "converted")
        output_name = os.path.splitext(output_name)[0] + ".docx"
        
        # Convert image to Word using OCR
        buffer = image_to_word(image_path)
        
        response = send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=output_name
        )
        
        @after_this_request
        def cleanup(response):
            safe_remove(image_path)
            return response
        
        return response
        
    except Exception as e:
        safe_remove(image_path)
        return jsonify({"error": f"Conversion failed: {str(e)}"}), 500

# -----------------------------------------------------------------------------
# NEW: PDF to Searchable PDF (OCR)
# -----------------------------------------------------------------------------
@app.route('/api/pdf-to-searchable', methods=['POST'])
def api_pdf_to_searchable():
    """Convert scanned PDF to searchable PDF with OCR text layer"""
    if not OCR_AVAILABLE:
        return jsonify({"error": "OCR functionality is not available."}), 500
    
    cleanup_temp()
    files = request.files.getlist('files')
    if not files or len(files) != 1:
        return jsonify({"error": "Upload exactly one PDF."}), 400
    
    paths = save_uploads(files)
    pdf_path = paths[0]
    
    if ext_of(pdf_path) not in ALLOWED_PDF_EXT:
        safe_remove(pdf_path)
        return jsonify({"error": "Only PDF files are allowed."}), 400

    try:
        original_name = secure_filename(files[0].filename)
        output_name = generate_unique_filename(original_name, "searchable")
        output_name = os.path.splitext(output_name)[0] + ".pdf"
        
        # Create searchable PDF
        searchable_path = ocr_pdf_to_searchable_pdf(pdf_path)
        
        with open(searchable_path, 'rb') as f:
            buffer = io.BytesIO(f.read())
        
        buffer.seek(0)
        
        response = send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=output_name
        )
        
        @after_this_request
        def cleanup(response):
            safe_remove(pdf_path)
            safe_remove(searchable_path)
            return response
        
        return response
        
    except Exception as e:
        safe_remove(pdf_path)
        return jsonify({"error": f"Conversion failed: {str(e)}"}), 500

# -----------------------------------------------------------------------------
# Existing APIs (unchanged but with OCR awareness)
# -----------------------------------------------------------------------------
@app.route('/api/merge-pdf', methods=['POST'])
def api_merge_pdf():
    cleanup_temp()
    files = request.files.getlist('files')
    if not files or len(files) < 2:
        return jsonify({"error": "Upload at least two PDFs."}), 400
    
    paths = save_uploads(files)
    for p in paths:
        if ext_of(p) not in ALLOWED_PDF_EXT:
            safe_remove_all(paths)
            return jsonify({"error": "Only PDF files are allowed."}), 400

    merger = PdfMerger()
    try:
        for p in paths:
            merger.append(p)
        
        original_name = secure_filename(files[0].filename)
        output_name = generate_unique_filename(original_name, "merged")
        output_name = os.path.splitext(output_name)[0] + ".pdf"
        
        buffer = io.BytesIO()
        merger.write(buffer)
        buffer.seek(0)
        
        response = send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=output_name
        )
        
        @after_this_request
        def cleanup(response):
            safe_remove_all(paths)
            merger.close()
            return response
        
        return response
        
    except Exception as e:
        safe_remove_all(paths)
        merger.close()
        return jsonify({"error": f"Merging failed: {str(e)}"}), 500

@app.route('/api/split-pdf', methods=['POST'])
def api_split_pdf():
    cleanup_temp()
    files = request.files.getlist('files')
    if not files or len(files) != 1:
        return jsonify({"error": "Upload exactly one PDF."}), 400
    
    ranges_str = request.form.get('ranges', '').strip()
    if not ranges_str:
        return jsonify({"error": "Page ranges are required."}), 400
    
    paths = save_uploads(files)
    pdf_path = paths[0]
    if ext_of(pdf_path) not in ALLOWED_PDF_EXT:
        safe_remove(pdf_path)
        return jsonify({"error": "Only PDF files are allowed."}), 400
    
    try:
        reader = PdfReader(pdf_path)
        total_pages = len(reader.pages)
        
        ranges = []
        parts = [p.strip() for p in ranges_str.split(',') if p.strip()]
        for part in parts:
            if '-' in part:
                start, end = part.split('-', 1)
                try:
                    start = int(start); end = int(end)
                    if 1 <= start <= total_pages and 1 <= end <= total_pages:
                        ranges.append((min(start, end)-1, max(start, end)))
                    else:
                        safe_remove(pdf_path)
                        return jsonify({"error": f"Page range out of bounds (1-{total_pages})."}), 400
                except ValueError:
                    safe_remove(pdf_path)
                    return jsonify({"error": "Invalid page range format."}), 400
            else:
                try:
                    page = int(part)
                    if 1 <= page <= total_pages:
                        ranges.append((page-1, page))
                    else:
                        safe_remove(pdf_path)
                        return jsonify({"error": f"Page out of bounds (1-{total_pages})."}), 400
                except ValueError:
                    safe_remove(pdf_path)
                    return jsonify({"error": "Invalid page number."}), 400
        
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for i, (start_idx, end_page) in enumerate(ranges):
                writer = PdfWriter()
                for page_idx in range(start_idx, end_page):
                    writer.add_page(reader.pages[page_idx])
                
                split_buffer = io.BytesIO()
                writer.write(split_buffer)
                split_buffer.seek(0)
                
                original_name = secure_filename(files[0].filename)
                split_name = generate_unique_filename(original_name, f"split_{i+1}")
                split_name = os.path.splitext(split_name)[0] + ".pdf"
                
                zipf.writestr(split_name, split_buffer.getvalue())
                writer.close()
        
        zip_buffer.seek(0)
        original_name = secure_filename(files[0].filename)
        zip_name = generate_unique_filename(original_name, "split_parts")
        zip_name = os.path.splitext(zip_name)[0] + ".zip"
        
        response = send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name=zip_name
        )
        
        @after_this_request
        def cleanup(response):
            safe_remove(pdf_path)
            return response
        
        return response
        
    except Exception as e:
        safe_remove(pdf_path)
        return jsonify({"error": f"Splitting failed: {str(e)}"}), 500

@app.route('/api/delete-pages-pdf', methods=['POST'])
def api_delete_pages_pdf():
    cleanup_temp()
    pages_str = request.form.get('pages', '').strip()
    files = request.files.getlist('files')
    if not files or len(files) != 1:
        return jsonify({"error": "Upload exactly one PDF."}), 400
    if not pages_str:
        return jsonify({"error": "Pages to remove are required."}), 400
    
    paths = save_uploads(files)
    pdf_path = paths[0]
    if ext_of(pdf_path) not in ALLOWED_PDF_EXT:
        safe_remove(pdf_path)
        return jsonify({"error": "Only PDF files are allowed."}), 400

    pages_to_remove = parse_pages(pages_str)
    writer = PdfWriter()
    
    try:
        reader = PdfReader(pdf_path)
        total = len(reader.pages)
        
        original_name = secure_filename(files[0].filename)
        output_name = generate_unique_filename(original_name, "pages_removed")
        output_name = os.path.splitext(output_name)[0] + ".pdf"
        
        for i in range(total):
            if (i+1) not in pages_to_remove:
                writer.add_page(reader.pages[i])
        
        buffer = io.BytesIO()
        writer.write(buffer)
        buffer.seek(0)
        
        response = send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=output_name
        )
        
        @after_this_request
        def cleanup(response):
            safe_remove(pdf_path)
            writer.close()
            return response
        
        return response
        
    except Exception as e:
        safe_remove(pdf_path)
        writer.close()
        return jsonify({"error": f"Page removal failed: {str(e)}"}), 500

@app.route('/api/rotate-pdf', methods=['POST'])
def api_rotate_pdf():
    cleanup_temp()
    rotation = int(request.form.get('rotation', '90'))
    files = request.files.getlist('files')
    if not files or len(files) != 1:
        return jsonify({"error": "Upload exactly one PDF."}), 400
    
    paths = save_uploads(files)
    pdf_path = paths[0]
    if ext_of(pdf_path) not in ALLOWED_PDF_EXT:
        safe_remove(pdf_path)
        return jsonify({"error": "Only PDF files are allowed."}), 400

    writer = PdfWriter()
    try:
        reader = PdfReader(pdf_path)
        
        original_name = secure_filename(files[0].filename)
        output_name = generate_unique_filename(original_name, f"rotated_{rotation}")
        output_name = os.path.splitext(output_name)[0] + ".pdf"
        
        for page in reader.pages:
            page.rotate(rotation)
            writer.add_page(page)
        
        buffer = io.BytesIO()
        writer.write(buffer)
        buffer.seek(0)
        
        response = send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=output_name
        )
        
        @after_this_request
        def cleanup(response):
            safe_remove(pdf_path)
            writer.close()
            return response
        
        return response
        
    except Exception as e:
        safe_remove(pdf_path)
        writer.close()
        return jsonify({"error": f"Rotation failed: {str(e)}"}), 500

@app.route('/api/lock-pdf', methods=['POST'])
def api_lock_pdf():
    cleanup_temp()
    pin = request.form.get('pin', '').strip()
    if not pin or len(pin) != 4 or not pin.isdigit():
        return jsonify({"error": "PIN must be exactly 4 digits."}), 400
    
    files = request.files.getlist('files')
    if not files or len(files) != 1:
        return jsonify({"error": "Upload exactly one PDF."}), 400
    
    paths = save_uploads(files)
    pdf_path = paths[0]
    if ext_of(pdf_path) not in ALLOWED_PDF_EXT:
        safe_remove(pdf_path)
        return jsonify({"error": "Only PDF files are allowed."}), 400

    writer = PdfWriter()
    try:
        reader = PdfReader(pdf_path)
        
        original_name = secure_filename(files[0].filename)
        output_name = generate_unique_filename(original_name, "locked")
        output_name = os.path.splitext(output_name)[0] + ".pdf"
        
        for page in reader.pages:
            writer.add_page(page)
        writer.encrypt(pin)
        
        buffer = io.BytesIO()
        writer.write(buffer)
        buffer.seek(0)
        
        response = send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=output_name
        )
        
        @after_this_request
        def cleanup(response):
            safe_remove(pdf_path)
            writer.close()
            return response
        
        return response
        
    except Exception as e:
        safe_remove(pdf_path)
        writer.close()
        return jsonify({"error": f"Locking failed: {str(e)}"}), 500

@app.route('/api/unlock-pdf', methods=['POST'])
def api_unlock_pdf():
    cleanup_temp()
    password = request.form.get('password', '').strip()
    files = request.files.getlist('files')
    if not files or len(files) != 1:
        return jsonify({"error": "Upload exactly one PDF."}), 400
    if not password:
        return jsonify({"error": "Password is required."}), 400
    
    paths = save_uploads(files)
    pdf_path = paths[0]
    if ext_of(pdf_path) not in ALLOWED_PDF_EXT:
        safe_remove(pdf_path)
        return jsonify({"error": "Only PDF files are allowed."}), 400

    writer = PdfWriter()
    try:
        reader = PdfReader(pdf_path)
        
        original_name = secure_filename(files[0].filename)
        output_name = generate_unique_filename(original_name, "unlocked")
        output_name = os.path.splitext(output_name)[0] + ".pdf"
        
        if reader.is_encrypted:
            if not reader.decrypt(password):
                safe_remove(pdf_path)
                writer.close()
                return jsonify({"error": "Incorrect password."}), 400
        for page in reader.pages:
            writer.add_page(page)
        
        buffer = io.BytesIO()
        writer.write(buffer)
        buffer.seek(0)
        
        response = send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=output_name
        )
        
        @after_this_request
        def cleanup(response):
            safe_remove(pdf_path)
            writer.close()
            return response
        
        return response
        
    except Exception as e:
        safe_remove(pdf_path)
        writer.close()
        return jsonify({"error": f"Unlocking failed: {str(e)}"}), 500

@app.route('/api/text-to-pdf', methods=['POST'])
def api_text_to_pdf():
    cleanup_temp()
    text = (request.form.get('text') or '').strip()
    if not text:
        return jsonify({"error": "Text content is required."}), 400

    cleaned_text = clean_text_for_xml(text)
    
    unique_id = str(uuid.uuid4())[:12]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_name = f"text_converted_{timestamp}_{unique_id}.pdf"

    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    left_margin = 50
    top = height - 50
    line_height = 14
    
    lines = cleaned_text.splitlines()
    for line in lines[:500]:
        if line.strip():
            for chunk in wrap_text(line, max_chars=95):
                c.drawString(left_margin, top, chunk)
                top -= line_height
                if top < 50:
                    c.showPage()
                    top = height - 50
        else:
            top -= line_height
            if top < 50:
                c.showPage()
                top = height - 50
    
    c.save()
    buffer.seek(0)
    
    return send_file(
        buffer,
        mimetype='application/pdf',
        as_attachment=True,
        download_name=output_name
    )

@app.route('/api/text-to-word', methods=['POST'])
def api_text_to_word():
    cleanup_temp()
    text = (request.form.get('text') or '').strip()
    if not text:
        return jsonify({"error": "Text content is required."}), 400
    
    cleaned_text = clean_text_for_xml(text)
    
    unique_id = str(uuid.uuid4())[:12]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_name = f"text_converted_{timestamp}_{unique_id}.docx"
    
    doc = Document()
    
    if cleaned_text:
        lines = cleaned_text.splitlines()
        for line in lines[:500]:
            if line.strip():
                safe_add_paragraph(doc, line)
    else:
        doc.add_paragraph("No text content provided.")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return send_file(
        buffer,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=output_name
    )

@app.route('/api/images-to-pdf', methods=['POST'])
def api_images_to_pdf():
    cleanup_temp()
    files = request.files.getlist('files')
    if not files or len(files) < 1:
        return jsonify({"error": "Upload at least one image."}), 400
    
    paths = save_uploads(files)
    for p in paths:
        if ext_of(p) not in ALLOWED_IMAGE_EXT:
            safe_remove_all(paths)
            return jsonify({"error": "Only image files are allowed."}), 400

    try:
        original_name = secure_filename(files[0].filename)
        output_name = generate_unique_filename(original_name, "images_to_pdf")
        output_name = os.path.splitext(output_name)[0] + ".pdf"
        
        images = []
        for p in paths:
            img = Image.open(p).convert('RGB')
            images.append(img)
        
        buffer = io.BytesIO()
        if len(images) == 1:
            images[0].save(buffer, format='PDF', save_all=True)
        else:
            first, rest = images[0], images[1:]
            first.save(buffer, format='PDF', save_all=True, append_images=rest)
        
        buffer.seek(0)
        
        response = send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=output_name
        )
        
        @after_this_request
        def cleanup(response):
            safe_remove_all(paths)
            return response
        
        return response
        
    except Exception as e:
        safe_remove_all(paths)
        return jsonify({"error": f"Conversion failed: {str(e)}"}), 500

# -----------------------------------------------------------------------------
# Health check and error handlers
# -----------------------------------------------------------------------------
@app.route('/health', methods=['GET'])
def health_check():
    return jsonify({
        "status": "healthy", 
        "timestamp": datetime.now().isoformat(),
        "ocr_available": OCR_AVAILABLE,
        "ocr_enabled": OCR_ENABLED
    }), 200

@app.errorhandler(404)
def page_not_found(e):
    return jsonify({"error": "Page not found. Please check the URL."}), 404

@app.errorhandler(413)
def too_large(e):
    return jsonify({"error": "File too large (max 50 MB)."}), 413

@app.errorhandler(400)
def bad_request(e):
    return jsonify({"error": str(e.description) if e.description else "Bad request."}), 400

@app.errorhandler(500)
def server_error(e):
    return jsonify({"error": "Internal server error."}), 500

@app.route('/static/<path:filename>')
def serve_static(filename):
    return send_from_directory('static', filename)

@app.after_request
def after_request(response):
    response.headers.add('Access-Control-Allow-Origin', '*')
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
    response.headers.add('Access-Control-Allow-Methods', 'GET,PUT,POST,DELETE,OPTIONS')
    return response

def cleanup_cache():
    while True:
        time.sleep(300)
        try:
            if len(conversion_cache) > 100:
                keys = list(conversion_cache.keys())[:-100]
                for key in keys:
                    del conversion_cache[key]
        except:
            pass

import threading
cache_cleaner = threading.Thread(target=cleanup_cache, daemon=True)
cache_cleaner.start()

if __name__ == '__main__':
    print(f"Starting iMasterPDF with {MAX_WORKERS} workers")
    print(f"Cache enabled: {CACHE_ENABLED}")
    print(f"OCR enabled: {OCR_ENABLED}")
    print(f"OCR available: {OCR_AVAILABLE}")
    app.run(host='0.0.0.0', port=8000, debug=False, threaded=True)