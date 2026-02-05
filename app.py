import os
import io
import shutil
import tempfile
import uuid
import re
import time
import hashlib
import zipfile
from datetime import datetime, timedelta
from concurrent.futures import ThreadPoolExecutor
from functools import lru_cache
import concurrent.futures
import threading
import json

from flask import Flask, render_template, send_file, request, abort, Response, jsonify, send_from_directory, after_this_request
from werkzeug.utils import secure_filename
from werkzeug.middleware.proxy_fix import ProxyFix

from PyPDF2 import PdfReader, PdfWriter, PdfMerger
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4, letter
from PIL import Image, ImageOps, ImageEnhance, ImageFilter

# OCR Libraries - Essential for handling image PDFs
try:
    import pytesseract
    from pdf2image import convert_from_path, convert_from_bytes
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# -----------------------------------------------------------------------------
# Flask app configuration
# -----------------------------------------------------------------------------
app = Flask(__name__)
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1, x_port=1)

# Performance settings
MAX_CONTENT_LENGTH = 100 * 1024 * 1024  # 100 MB per file
UPLOAD_DIR = os.path.join(tempfile.gettempdir(), "imasterpdf_uploads")
OUTPUT_DIR = os.path.join(tempfile.gettempdir(), "imasterpdf_outputs")
CLEANUP_AGE_MINUTES = 30
MAX_WORKERS = 6
MAX_PAGES_TO_EXTRACT = 200
CACHE_ENABLED = True
OCR_ENABLED = OCR_AVAILABLE  # Enable OCR if available

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'dev-secret-key-change-in-production')

# Thread pools for parallel processing
io_executor = ThreadPoolExecutor(max_workers=MAX_WORKERS)

# Cache for repeated conversions with TTL
conversion_cache = {}
CACHE_TTL_SECONDS = 3600  # 1 hour

ALLOWED_IMAGE_EXT = {'.jpg', '.jpeg', '.png', '.webp', '.bmp', '.tiff', '.tif', '.gif'}
ALLOWED_PDF_EXT = {'.pdf'}
ALLOWED_WORD_EXT = {'.docx', '.doc'}
ALLOWED_TEXT_EXT = {'.txt'}

# -----------------------------------------------------------------------------
# PDF Processing with OCR Support
# -----------------------------------------------------------------------------
class PDFProcessor:
    """Handles both text-based and image-based PDFs with OCR support"""
    
    @staticmethod
    def extract_text_from_pdf(pdf_path, use_ocr=False, languages=['eng']):
        """
        Extract text from PDF with OCR fallback
        Args:
            pdf_path: Path to PDF file
            use_ocr: Force OCR even if text is detected
            languages: List of language codes for OCR
        Returns:
            Extracted text
        """
        start_time = time.time()
        
        # Check cache first
        if CACHE_ENABLED:
            file_hash = hashlib.md5(pdf_path.encode()).hexdigest()
            file_hash += f"_{use_ocr}_{'_'.join(languages)}"
            if file_hash in conversion_cache:
                cache_time, text = conversion_cache[file_hash]
                if time.time() - cache_time < CACHE_TTL_SECONDS:
                    return text
        
        # Step 1: Try standard text extraction first
        text = PDFProcessor._extract_standard_text(pdf_path)
        
        # Step 2: If text is insufficient or OCR requested, use OCR
        if (not text or len(text.strip()) < 100 or use_ocr) and OCR_ENABLED:
            ocr_text = PDFProcessor._extract_ocr_text(pdf_path, languages)
            if ocr_text and len(ocr_text.strip()) > len(text.strip()):
                text = ocr_text
        
        if CACHE_ENABLED:
            conversion_cache[file_hash] = (time.time(), text)
        
        print(f"Text extraction completed in {time.time() - start_time:.2f}s")
        return text
    
    @staticmethod
    def _extract_standard_text(pdf_path):
        """Extract text from text-based PDFs"""
        try:
            # Method 1: PyPDF2 extraction
            with open(pdf_path, 'rb') as file:
                reader = PdfReader(file)
                text_parts = []
                
                for i, page in enumerate(reader.pages[:MAX_PAGES_TO_EXTRACT]):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_parts.append(page_text.strip())
                    except:
                        continue
                
                if text_parts:
                    return "\n\n".join(text_parts)
                
            return ""
                
        except Exception as e:
            print(f"Standard text extraction failed: {e}")
            return ""
    
    @staticmethod
    def _extract_ocr_text(pdf_path, languages=['eng']):
        """Extract text from scanned/image PDFs using OCR"""
        if not OCR_ENABLED:
            print("OCR not available")
            return ""
        
        try:
            # Convert PDF to images
            images = PDFProcessor._pdf_to_images(pdf_path)
            if not images:
                print("No images extracted from PDF")
                return ""
            
            # Process images in parallel
            texts = PDFProcessor._parallel_ocr_processing(images, languages)
            
            # Combine results
            combined_text = "\n\n".join(texts)
            return combined_text.strip()
            
        except Exception as e:
            print(f"OCR extraction failed: {e}")
            return ""
    
    @staticmethod
    def _pdf_to_images(pdf_path, max_pages=50):
        """Convert PDF pages to images"""
        try:
            images = convert_from_bytes(
                open(pdf_path, 'rb').read(),
                dpi=300,  # Good balance for OCR
                thread_count=2,
                fmt='jpeg',
                first_page=1,
                last_page=min(max_pages, MAX_PAGES_TO_EXTRACT),
                grayscale=True  # Grayscale for better OCR
            )
            return images
        except Exception as e:
            print(f"PDF to image conversion failed: {e}")
            # Try alternative method
            try:
                images = convert_from_path(
                    pdf_path,
                    dpi=300,
                    thread_count=1,
                    fmt='jpeg',
                    first_page=1,
                    last_page=min(max_pages, MAX_PAGES_TO_EXTRACT)
                )
                return images
            except:
                return []
    
    @staticmethod
    def _parallel_ocr_processing(images, languages):
        """Process images with OCR in parallel"""
        def process_single_image(img):
            try:
                # Preprocess image for better OCR
                img = ImageOps.exif_transpose(img)
                
                # Convert to grayscale if not already
                if img.mode != 'L':
                    img = img.convert('L')
                
                # Enhance contrast
                enhancer = ImageEnhance.Contrast(img)
                img = enhancer.enhance(1.5)
                
                # Apply sharpening
                img = img.filter(ImageFilter.SHARPEN)
                
                # Perform OCR
                text = pytesseract.image_to_string(
                    img,
                    lang='+'.join(languages),
                    config='--psm 3 --oem 3 -c preserve_interword_spaces=1'
                )
                return text.strip()
            except Exception as e:
                print(f"Single image OCR failed: {e}")
                return ""
        
        # Process in parallel
        with ThreadPoolExecutor(max_workers=min(4, len(images))) as executor:
            futures = [executor.submit(process_single_image, img) for img in images]
            results = []
            for future in concurrent.futures.as_completed(futures):
                try:
                    text = future.result(timeout=30)
                    if text:
                        results.append(text)
                except:
                    pass
        
        return results
    
    @staticmethod
    def is_image_pdf(pdf_path, threshold=100):
        """
        Detect if PDF is image-based (scanned)
        Args:
            pdf_path: Path to PDF file
            threshold: Minimum characters to consider as text PDF
        Returns:
            True if image-based, False if text-based
        """
        try:
            # Try to extract text from first few pages
            with open(pdf_path, 'rb') as file:
                reader = PdfReader(file)
                text = ""
                
                for i, page in enumerate(reader.pages[:3]):
                    try:
                        page_text = page.extract_text() or ""
                        text += page_text
                    except:
                        pass
                
                # Check if we got meaningful text
                if len(text.strip()) >= threshold:
                    # Count alphabetic characters
                    alpha_chars = sum(1 for c in text if c.isalpha())
                    total_chars = len(text)
                    
                    if total_chars > 0 and alpha_chars / total_chars > 0.1:  # More than 10% alphabetic
                        return False  # Text PDF
                
                return True  # Image PDF
                
        except Exception as e:
            print(f"PDF detection failed: {e}")
            return True  # Assume image PDF if detection fails
    
    @staticmethod
    def create_searchable_pdf(pdf_path, output_path, languages=['eng']):
        """
        Create a searchable PDF from scanned PDF
        Adds invisible text layer over images
        """
        if not OCR_ENABLED:
            raise Exception("OCR not available")
        
        try:
            # Convert PDF to images
            images = PDFProcessor._pdf_to_images(pdf_path)
            if not images:
                raise Exception("No images extracted")
            
            # Create new PDF with text layer
            c = canvas.Canvas(output_path, pagesize=A4)
            width, height = A4
            
            for img in images:
                # Convert image to bytes
                img_buffer = io.BytesIO()
                img.save(img_buffer, format='JPEG', quality=85)
                img_buffer.seek(0)
                
                # Add image as background
                c.drawImage(img_buffer, 0, 0, width=width, height=height)
                
                # Extract text from image
                text = pytesseract.image_to_string(img, lang='+'.join(languages))
                
                if text.strip():
                    # Add invisible text layer
                    c.setFont("Helvetica", 1)  # Tiny font
                    c.setFillColorRGB(1, 1, 1, alpha=0)  # Fully transparent
                    
                    # Add text at very small position (invisible but searchable)
                    c.drawString(1, 1, text[:1000])  # Limit text
                
                c.showPage()
            
            c.save()
            return True
            
        except Exception as e:
            print(f"Searchable PDF creation failed: {e}")
            return False

# -----------------------------------------------------------------------------
# Performance optimization utilities
# -----------------------------------------------------------------------------
class UltraFastProcessor:
    """Optimized processor for ultra-fast conversions"""
    
    @staticmethod
    def clean_text_for_xml(text):
        """Ultra-fast text cleaning with regex compilation"""
        if not text:
            return ""
        
        # Pre-compiled regex patterns
        control_chars = re.compile(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]')
        
        # Fast operations
        text = text.replace('\x00', '')
        text = control_chars.sub('', text)
        
        # Fast Unicode replacements
        replacements = [
            ('\u2028', ' '),
            ('\u2029', ' '),
            ('\uFEFF', ''),
        ]
        
        for old, new in replacements:
            text = text.replace(old, new)
        
        return text
    
    @staticmethod
    def fast_extract_text(pdf_path, use_ocr=False, languages=['eng']):
        """Ultra-fast text extraction with intelligent fallback and OCR support"""
        start_time = time.time()
        
        # Check cache first
        if CACHE_ENABLED:
            file_hash = hashlib.md5(pdf_path.encode()).hexdigest() + "_" + str(use_ocr) + "_" + "_".join(languages)
            if file_hash in conversion_cache:
                cache_time, text = conversion_cache[file_hash]
                if time.time() - cache_time < CACHE_TTL_SECONDS:
                    return text
        
        # Determine extraction strategy
        file_size = os.path.getsize(pdf_path)
        
        # Strategy 1: Try regular extraction first (for text-based PDFs)
        try:
            with open(pdf_path, 'rb') as file:
                reader = PdfReader(file)
                text = []
                for i, page in enumerate(reader.pages[:MAX_PAGES_TO_EXTRACT]):
                    page_text = page.extract_text()
                    if page_text and len(page_text.strip()) > 50:  # Check if meaningful text
                        text.append(page_text)
                
                if text and len("".join(text).strip()) > 100:  # If enough text found
                    result = "\n".join(text)
                    if CACHE_ENABLED:
                        conversion_cache[file_hash] = (time.time(), result)
                    return result
        except:
            pass
        
        # Strategy 2: If no/insufficient text or OCR requested, use OCR
        if OCR_ENABLED and (use_ocr or file_size < 50 * 1024 * 1024):  # OCR for < 50MB files
            try:
                # Use PDFProcessor OCR extraction
                result = PDFProcessor._extract_ocr_text(pdf_path, languages)
                if result and len(result.strip()) > 50:
                    if CACHE_ENABLED:
                        conversion_cache[file_hash] = (time.time(), result)
                    return result
            except Exception as e:
                print(f"OCR extraction attempt failed: {e}")
        
        # Strategy 3: Fallback to minimal extraction
        try:
            with open(pdf_path, 'rb') as file:
                reader = PdfReader(file)
                text_parts = []
                for i, page in enumerate(reader.pages[:10]):  # Limited pages
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                    except:
                        continue
                
                result = "\n".join(text_parts) if text_parts else ""
                if CACHE_ENABLED:
                    conversion_cache[file_hash] = (time.time(), result)
                return result
        except:
            return ""
    
    @staticmethod
    def is_image_pdf(pdf_path):
        """Detect if PDF is image-based (scanned) - optimized version"""
        return PDFProcessor.is_image_pdf(pdf_path)

# -----------------------------------------------------------------------------
# Utility helpers
# -----------------------------------------------------------------------------
def ext_of(filename):
    return os.path.splitext(filename.lower())[1]

def validate_file(stream):
    stream.seek(0, os.SEEK_END)
    size = stream.tell()
    stream.seek(0)
    if size < 1024:
        abort(Response("File too small (min 1 KB).", status=400))
    if size > MAX_CONTENT_LENGTH:
        abort(Response(f"File too large (max {MAX_CONTENT_LENGTH // (1024*1024)} MB).", status=400))

def generate_unique_filename(original_filename, suffix=""):
    """Generate a unique filename with UUID and timestamp"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    unique_id = str(uuid.uuid4())[:12]
    name, ext = os.path.splitext(original_filename)
    safe_name = secure_filename(name)
    
    if suffix:
        return f"{safe_name}_{suffix}_{timestamp}_{unique_id}{ext}"
    return f"{safe_name}_{timestamp}_{unique_id}{ext}"

def save_uploads(files):
    saved = []
    for storage in files:
        validate_file(storage.stream)
        filename = secure_filename(storage.filename)
        if not filename:
            abort(Response("Invalid filename.", status=400))
        
        unique_filename = generate_unique_filename(storage.filename)
        path = os.path.join(UPLOAD_DIR, unique_filename)
        storage.save(path)
        saved.append(path)
    return saved

def cleanup_temp():
    """Fast cleanup with bulk operations"""
    cutoff = datetime.utcnow() - timedelta(minutes=CLEANUP_AGE_MINUTES)
    
    for base in (UPLOAD_DIR, OUTPUT_DIR):
        try:
            for name in os.listdir(base):
                path = os.path.join(base, name)
                try:
                    mtime = datetime.utcfromtimestamp(os.path.getmtime(path))
                    if mtime < cutoff:
                        if os.path.isdir(path):
                            shutil.rmtree(path, ignore_errors=True)
                        else:
                            os.remove(path)
                except:
                    pass
        except:
            pass

def safe_remove(path):
    try:
        if os.path.exists(path):
            os.remove(path)
    except:
        pass

def safe_remove_all(paths):
    for path in paths:
        safe_remove(path)

def clean_text_for_xml(text):
    """Clean text for XML/Word document safety"""
    return UltraFastProcessor.clean_text_for_xml(text)

def safe_add_paragraph(doc, text):
    """Safely add a paragraph to a Word document"""
    try:
        cleaned_text = clean_text_for_xml(text)
        if cleaned_text.strip():
            doc.add_paragraph(cleaned_text.strip())
    except:
        pass

def parse_pages(pages_str):
    """Fast page parsing with set operations"""
    pages = set()
    if not pages_str:
        return pages
    
    parts = [p.strip() for p in pages_str.split(',') if p.strip()]
    for part in parts:
        if '-' in part:
            try:
                a, b = map(int, part.split('-', 1))
                pages.update(range(min(a, b), max(a, b) + 1))
            except:
                abort(Response("Invalid page range.", status=400))
        else:
            try:
                pages.add(int(part))
            except:
                abort(Response("Invalid page number.", status=400))
    return pages

def wrap_text(text, max_chars=95):
    """Fast text wrapping"""
    if len(text) <= max_chars:
        return [text]
    
    words = text.split(' ')
    lines = []
    current_line = []
    current_length = 0
    
    for word in words:
        word_length = len(word)
        if current_length + word_length + (1 if current_line else 0) <= max_chars:
            current_line.append(word)
            current_length += word_length + (1 if current_line else 0)
        else:
            if current_line:
                lines.append(' '.join(current_line))
            current_line = [word]
            current_length = word_length
    
    if current_line:
        lines.append(' '.join(current_line))
    
    return lines

# -----------------------------------------------------------------------------
# SPA Routes for each tool page
# -----------------------------------------------------------------------------

@app.route('/')
@app.route('/index')
@app.route('/index.html')
def index():
    return render_template('index.html')

@app.route('/split')
@app.route('/split.html')
def split_pdf():
    return render_template('split.html')

@app.route('/mergepdf')
@app.route('/mergepdf.html')
def merge_pdf():
    return render_template('mergepdf.html')

@app.route('/deletepdf')
@app.route('/deletepdf.html')
def delete_pdf():
    return render_template('deletepdf.html')

@app.route('/rotatepdf')
@app.route('/rotatepdf.html')
def rotate_pdf():
    return render_template('rotatepdf.html')

@app.route('/pdftoword')
@app.route('/pdftoword.html')
def pdf_to_word():
    return render_template('pdftoword.html')

@app.route('/lockpdf')
@app.route('/lockpdf.html')
def lock_pdf():
    return render_template('lockpdf.html')

@app.route('/unlockpdf')
@app.route('/unlockpdf.html')
def unlock_pdf():
    return render_template('unlockpdf.html')

@app.route('/wordtopdf')
@app.route('/wordtopdf.html')
def word_to_pdf():
    return render_template('wordtopdf.html')

@app.route('/mergeword')
@app.route('/mergeword.html')
def merge_word():
    return render_template('mergeword.html')

@app.route('/wordtotext')
@app.route('/wordtotext.html')
def word_to_text():
    return render_template('wordtotext.html')

@app.route('/texttopdf')
@app.route('/texttopdf.html')
def text_to_pdf():
    return render_template('texttopdf.html')

@app.route('/texttoword')
@app.route('/texttoword.html')
def text_to_word():
    return render_template('texttoword.html')

@app.route('/imagestopdf')
@app.route('/imagestopdf.html')
def images_to_pdf():
    return render_template('imagestopdf.html')

# OCR-specific route
@app.route('/ocrpdf')
@app.route('/ocrpdf.html')
def ocr_pdf():
    return render_template('ocrpdf.html')

# -----------------------------------------------------------------------------
# Contact API
# -----------------------------------------------------------------------------
@app.route('/api/contact', methods=['POST'])
def api_contact():
    data = request.get_json(silent=True) or {}
    name = (data.get('name') or '').strip()
    email = (data.get('email') or '').strip()
    message = (data.get('message') or '').strip()
    if not name or not email or not message:
        return jsonify({"error": "Please provide name, email, and message."}), 400
    return jsonify({"status": "ok", "received": {"name": name, "email": email}}), 200

# -----------------------------------------------------------------------------
# PDF to Word API with OCR support
# -----------------------------------------------------------------------------

@app.route('/api/pdf-to-word', methods=['POST'])
def api_pdf_to_word():
    """Convert PDF to Word with automatic OCR detection"""
    start_time = time.time()
    cleanup_temp()
    
    files = request.files.getlist('files')
    if not files or len(files) != 1:
        return jsonify({"error": "Upload exactly one PDF."}), 400
    
    # Check if OCR is forced
    force_ocr = request.form.get('force_ocr', 'false').lower() == 'true'
    language = request.form.get('language', 'eng').strip()
    
    paths = save_uploads(files)
    pdf_path = paths[0]
    
    if ext_of(pdf_path) not in ALLOWED_PDF_EXT:
        safe_remove(pdf_path)
        return jsonify({"error": "Only PDF files are allowed."}), 400
    
    try:
        # Determine if we need OCR
        needs_ocr = force_ocr or PDFProcessor.is_image_pdf(pdf_path)
        
        if needs_ocr and not OCR_ENABLED:
            return jsonify({"error": "OCR is required for this PDF but OCR is not available. Please install OCR dependencies."}), 400
        
        # Extract text with appropriate method
        languages = [language] if language != 'eng' else ['eng']
        
        # Use UltraFastProcessor for extraction
        text = UltraFastProcessor.fast_extract_text(
            pdf_path, 
            use_ocr=needs_ocr, 
            languages=languages
        )
        
        if not text or len(text.strip()) < 10:
            return jsonify({"error": "Could not extract any text from the PDF."}), 400
        
        # Create Word document
        doc = Document()
        
        # Clean and add text
        cleaned_text = clean_text_for_xml(text)
        paragraphs = [p.strip() for p in cleaned_text.split('\n\n') if p.strip()]
        
        # Limit for very large documents
        if len(paragraphs) > 500:
            paragraphs = paragraphs[:500]
            doc.add_paragraph("[Document truncated - showing first 500 paragraphs]")
        
        # Add paragraphs
        for para in paragraphs:
            safe_add_paragraph(doc, para)
        
        # Add metadata about conversion
        if needs_ocr:
            doc.add_paragraph(f"\n[Converted using OCR - Language: {language}]")
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Generate output filename
        original_name = secure_filename(files[0].filename)
        suffix = "ocr_converted" if needs_ocr else "converted"
        output_name = generate_unique_filename(original_name, suffix)
        output_name = os.path.splitext(output_name)[0] + ".docx"
        
        response = send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=output_name
        )
        
        @after_this_request
        def cleanup(response):
            safe_remove(pdf_path)
            return response
        
        print(f"PDF to Word conversion completed in {time.time() - start_time:.2f}s (OCR: {needs_ocr})")
        return response
        
    except Exception as e:
        safe_remove(pdf_path)
        return jsonify({"error": f"Conversion failed: {str(e)}"}), 500

# -----------------------------------------------------------------------------
# OCR PDF API - Convert scanned PDF to searchable PDF
# -----------------------------------------------------------------------------

@app.route('/api/ocr-pdf', methods=['POST'])
def api_ocr_pdf():
    """Convert scanned/image PDF to searchable PDF with OCR"""
    if not OCR_ENABLED:
        return jsonify({"error": "OCR is not available. Please install required packages."}), 400
    
    start_time = time.time()
    cleanup_temp()
    
    files = request.files.getlist('files')
    if not files or len(files) != 1:
        return jsonify({"error": "Upload exactly one PDF."}), 400
    
    language = request.form.get('language', 'eng').strip()
    output_format = request.form.get('format', 'pdf').strip().lower()
    
    paths = save_uploads(files)
    pdf_path = paths[0]
    
    if ext_of(pdf_path) not in ALLOWED_PDF_EXT:
        safe_remove(pdf_path)
        return jsonify({"error": "Only PDF files are allowed."}), 400
    
    try:
        # Check if it's already a text PDF
        if not PDFProcessor.is_image_pdf(pdf_path):
            # It's already text-based, just return as-is or with extracted text
            if output_format == 'pdf':
                # Return original PDF
                with open(pdf_path, 'rb') as f:
                    buffer = io.BytesIO(f.read())
                buffer.seek(0)
                
                output_name = generate_unique_filename(files[0].filename, "already_searchable")
                output_name = os.path.splitext(output_name)[0] + ".pdf"
                
                response = send_file(
                    buffer,
                    mimetype='application/pdf',
                    as_attachment=True,
                    download_name=output_name
                )
                
                @after_this_request
                def cleanup(response):
                    safe_remove(pdf_path)
                    return response
                
                return response
            else:
                # Extract text from text PDF
                text = UltraFastProcessor.fast_extract_text(pdf_path, use_ocr=False)
        
        else:
            # It's an image PDF, process with OCR
            if output_format == 'pdf':
                # Create searchable PDF
                output_temp = os.path.join(tempfile.gettempdir(), f"ocr_{uuid.uuid4().hex}.pdf")
                
                success = PDFProcessor.create_searchable_pdf(
                    pdf_path, 
                    output_temp, 
                    languages=[language]
                )
                
                if not success:
                    safe_remove(pdf_path)
                    safe_remove(output_temp)
                    return jsonify({"error": "Failed to create searchable PDF."}), 500
                
                # Read the created PDF
                with open(output_temp, 'rb') as f:
                    buffer = io.BytesIO(f.read())
                buffer.seek(0)
                
                output_name = generate_unique_filename(files[0].filename, "searchable")
                output_name = os.path.splitext(output_name)[0] + ".pdf"
                
                response = send_file(
                    buffer,
                    mimetype='application/pdf',
                    as_attachment=True,
                    download_name=output_name
                )
                
                @after_this_request
                def cleanup(response):
                    safe_remove(pdf_path)
                    safe_remove(output_temp)
                    return response
                
                return response
                
            else:
                # Extract text using OCR
                text = PDFProcessor.extract_text_from_pdf(
                    pdf_path, 
                    use_ocr=True, 
                    languages=[language]
                )
        
        # Handle text-based outputs (Word or Text)
        if output_format == 'word':
            # Create Word document
            doc = Document()
            
            if text:
                paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
                for para in paragraphs[:300]:  # Limit
                    safe_add_paragraph(doc, para)
                
                doc.add_paragraph(f"\n[Extracted using OCR - Language: {language}]")
            else:
                doc.add_paragraph("No text could be extracted via OCR.")
            
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            output_name = generate_unique_filename(files[0].filename, "ocr_text")
            output_name = os.path.splitext(output_name)[0] + ".docx"
            
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            
        else:  # text format
            # Create plain text
            if not text:
                text = "No text could be extracted via OCR."
            
            buffer = io.BytesIO(text.encode('utf-8'))
            buffer.seek(0)
            
            output_name = generate_unique_filename(files[0].filename, "ocr_text")
            output_name = os.path.splitext(output_name)[0] + ".txt"
            
            mimetype = 'text/plain'
        
        response = send_file(
            buffer,
            mimetype=mimetype,
            as_attachment=True,
            download_name=output_name
        )
        
        @after_this_request
        def cleanup(response):
            safe_remove(pdf_path)
            return response
        
        print(f"OCR processing completed in {time.time() - start_time:.2f}s")
        return response
        
    except Exception as e:
        safe_remove(pdf_path)
        return jsonify({"error": f"OCR processing failed: {str(e)}"}), 500

# -----------------------------------------------------------------------------
# Extract Text from PDF API
# -----------------------------------------------------------------------------

@app.route('/api/extract-text', methods=['POST'])
def api_extract_text():
    """Extract text from PDF (supports both text and image PDFs)"""
    start_time = time.time()
    cleanup_temp()
    
    files = request.files.getlist('files')
    if not files or len(files) != 1:
        return jsonify({"error": "Upload exactly one PDF."}), 400
    
    use_ocr = request.form.get('use_ocr', 'auto').strip().lower()
    language = request.form.get('language', 'eng').strip()
    
    paths = save_uploads(files)
    pdf_path = paths[0]
    
    if ext_of(pdf_path) not in ALLOWED_PDF_EXT:
        safe_remove(pdf_path)
        return jsonify({"error": "Only PDF files are allowed."}), 400
    
    try:
        # Determine OCR strategy
        if use_ocr == 'force':
            force_ocr = True
        elif use_ocr == 'never':
            force_ocr = False
        else:  # 'auto'
            force_ocr = PDFProcessor.is_image_pdf(pdf_path)
        
        if force_ocr and not OCR_ENABLED:
            return jsonify({"error": "OCR is required but not available."}), 400
        
        # Extract text
        languages = [language] if language != 'eng' else ['eng']
        text = UltraFastProcessor.fast_extract_text(
            pdf_path, 
            use_ocr=force_ocr, 
            languages=languages
        )
        
        if not text or len(text.strip()) < 10:
            return jsonify({"error": "Could not extract any text from the PDF."}), 400
        
        # Create response
        output_name = generate_unique_filename(files[0].filename, "extracted_text")
        output_name = os.path.splitext(output_name)[0] + ".txt"
        
        # Add metadata header
        metadata = f"Extracted from: {files[0].filename}\n"
        metadata += f"Method: {'OCR' if force_ocr else 'Standard extraction'}\n"
        metadata += f"Language: {language}\n"
        metadata += f"Timestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        metadata += "=" * 50 + "\n\n"
        
        full_text = metadata + text
        
        buffer = io.BytesIO(full_text.encode('utf-8'))
        buffer.seek(0)
        
        response = send_file(
            buffer,
            mimetype='text/plain',
            as_attachment=True,
            download_name=output_name
        )
        
        @after_this_request
        def cleanup(response):
            safe_remove(pdf_path)
            return response
        
        print(f"Text extraction completed in {time.time() - start_time:.2f}s")
        return response
        
    except Exception as e:
        safe_remove(pdf_path)
        return jsonify({"error": f"Text extraction failed: {str(e)}"}), 500

# -----------------------------------------------------------------------------
# Detect PDF Type API
# -----------------------------------------------------------------------------

@app.route('/api/detect-pdf-type', methods=['POST'])
def api_detect_pdf_type():
    """Detect if PDF is text-based or image-based"""
    start_time = time.time()
    cleanup_temp()
    
    files = request.files.getlist('files')
    if not files or len(files) != 1:
        return jsonify({"error": "Upload exactly one PDF."}), 400
    
    paths = save_uploads(files)
    pdf_path = paths[0]
    
    if ext_of(pdf_path) not in ALLOWED_PDF_EXT:
        safe_remove(pdf_path)
        return jsonify({"error": "Only PDF files are allowed."}), 400
    
    try:
        is_image = PDFProcessor.is_image_pdf(pdf_path)
        
        # Try to extract some text for analysis
        sample_text = ""
        if not is_image:
            with open(pdf_path, 'rb') as file:
                reader = PdfReader(file)
                for i, page in enumerate(reader.pages[:2]):
                    try:
                        page_text = page.extract_text() or ""
                        sample_text += page_text[:500] + "\n"
                    except:
                        pass
        
        response_data = {
            "filename": files[0].filename,
            "is_image_pdf": is_image,
            "pdf_type": "scanned/image PDF" if is_image else "text-based PDF",
            "ocr_required": is_image,
            "sample_text": sample_text[:1000] if sample_text else "",
            "file_size": os.path.getsize(pdf_path),
            "ocr_available": OCR_ENABLED
        }
        
        @after_this_request
        def cleanup(response):
            safe_remove(pdf_path)
            return response
        
        print(f"PDF detection completed in {time.time() - start_time:.2f}s")
        return jsonify(response_data)
        
    except Exception as e:
        safe_remove(pdf_path)
        return jsonify({"error": f"PDF detection failed: {str(e)}"}), 500

# -----------------------------------------------------------------------------
# Tool APIs - PDF Operations (ULTRA-FAST) - ALL IMPLEMENTED
# -----------------------------------------------------------------------------

@app.route('/api/merge-pdf', methods=['POST'])
def api_merge_pdf():
    """Ultra-fast PDF merging - FULLY IMPLEMENTED"""
    start_time = time.time()
    cleanup_temp()
    
    files = request.files.getlist('files')
    if not files or len(files) < 2:
        return jsonify({"error": "Upload at least two PDFs."}), 400
    
    # Save files in parallel
    paths = save_uploads(files)
    
    # Validate all are PDFs
    for p in paths:
        if ext_of(p) not in ALLOWED_PDF_EXT:
            safe_remove_all(paths)
            return jsonify({"error": "Only PDF files are allowed."}), 400

    try:
        # Merge PDFs in memory
        merger = PdfMerger()
        for p in paths:
            merger.append(p, import_outline=False)  # Disable outline for speed
        
        # Generate output
        original_name = secure_filename(files[0].filename)
        output_name = generate_unique_filename(original_name, "merged")
        output_name = os.path.splitext(output_name)[0] + ".pdf"
        
        buffer = io.BytesIO()
        merger.write(buffer)
        buffer.seek(0)
        merger.close()
        
        # Prepare response
        response = send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=output_name
        )
        
        # Cleanup
        @after_this_request
        def cleanup(response):
            safe_remove_all(paths)
            return response
        
        print(f"Merged {len(files)} PDFs in {time.time() - start_time:.2f}s")
        return response
        
    except Exception as e:
        safe_remove_all(paths)
        return jsonify({"error": f"Merging failed: {str(e)}"}), 500

@app.route('/api/split-pdf', methods=['POST'])
def api_split_pdf():
    """Ultra-fast PDF splitting - FULLY IMPLEMENTED"""
    start_time = time.time()
    cleanup_temp()
    
    files = request.files.getlist('files')
    if not files or len(files) != 1:
        return jsonify({"error": "Upload exactly one PDF."}), 400
    
    ranges_str = request.form.get('ranges', '').strip()
    if not ranges_str:
        return jsonify({"error": "Page ranges are required."}), 400
    
    # Save file
    paths = save_uploads(files)
    pdf_path = paths[0]
    
    if ext_of(pdf_path) not in ALLOWED_PDF_EXT:
        safe_remove(pdf_path)
        return jsonify({"error": "Only PDF files are allowed."}), 400
    
    try:
        # Read PDF once
        with open(pdf_path, 'rb') as f:
            reader = PdfReader(f)
            total_pages = len(reader.pages)
        
        # Parse ranges
        ranges = []
        parts = [p.strip() for p in ranges_str.split(',') if p.strip()]
        for part in parts:
            if '-' in part:
                try:
                    start, end = map(int, part.split('-', 1))
                    if 1 <= start <= total_pages and 1 <= end <= total_pages:
                        ranges.append((min(start, end)-1, max(start, end)))
                    else:
                        raise ValueError
                except:
                    safe_remove(pdf_path)
                    return jsonify({"error": f"Page range out of bounds (1-{total_pages})."}), 400
            else:
                try:
                    page = int(part)
                    if 1 <= page <= total_pages:
                        ranges.append((page-1, page))
                    else:
                        raise ValueError
                except:
                    safe_remove(pdf_path)
                    return jsonify({"error": f"Page out of bounds (1-{total_pages})."}), 400
        
        # Create ZIP in memory with parallel processing
        zip_buffer = io.BytesIO()
        
        def create_split(range_idx, start_idx, end_page):
            writer = PdfWriter()
            for page_idx in range(start_idx, end_page):
                writer.add_page(reader.pages[page_idx])
            
            split_buffer = io.BytesIO()
            writer.write(split_buffer)
            writer.close()
            split_buffer.seek(0)
            
            original_name = secure_filename(files[0].filename)
            split_name = generate_unique_filename(original_name, f"split_{range_idx+1}")
            split_name = os.path.splitext(split_name)[0] + ".pdf"
            
            return split_name, split_buffer.getvalue()
        
        # Process splits in parallel
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
            with ThreadPoolExecutor(max_workers=4) as executor:
                futures = []
                for i, (start_idx, end_page) in enumerate(ranges):
                    future = executor.submit(create_split, i, start_idx, end_page)
                    futures.append(future)
                
                for future in concurrent.futures.as_completed(futures):
                    split_name, split_data = future.result()
                    zipf.writestr(split_name, split_data)
        
        zip_buffer.seek(0)
        original_name = secure_filename(files[0].filename)
        zip_name = generate_unique_filename(original_name, "split_parts")
        zip_name = os.path.splitext(zip_name)[0] + ".zip"
        
        response = send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name=zip_name
        )
        
        @after_this_request
        def cleanup(response):
            safe_remove(pdf_path)
            return response
        
        print(f"Split PDF into {len(ranges)} parts in {time.time() - start_time:.2f}s")
        return response
        
    except Exception as e:
        safe_remove(pdf_path)
        return jsonify({"error": f"Splitting failed: {str(e)}"}), 500

@app.route('/api/delete-pages-pdf', methods=['POST'])
def api_delete_pages_pdf():
    """Ultra-fast page deletion - FULLY IMPLEMENTED"""
    start_time = time.time()
    cleanup_temp()
    
    pages_str = request.form.get('pages', '').strip()
    files = request.files.getlist('files')
    if not files or len(files) != 1:
        return jsonify({"error": "Upload exactly one PDF."}), 400
    if not pages_str:
        return jsonify({"error": "Pages to remove are required."}), 400
    
    paths = save_uploads(files)
    pdf_path = paths[0]
    
    if ext_of(pdf_path) not in ALLOWED_PDF_EXT:
        safe_remove(pdf_path)
        return jsonify({"error": "Only PDF files are allowed."}), 400

    try:
        # Read and process in one pass
        with open(pdf_path, 'rb') as f:
            reader = PdfReader(f)
        
        pages_to_remove = parse_pages(pages_str)
        
        writer = PdfWriter()
        for i, page in enumerate(reader.pages):
            if (i + 1) not in pages_to_remove:
                writer.add_page(page)
        
        # Generate output
        original_name = secure_filename(files[0].filename)
        output_name = generate_unique_filename(original_name, "pages_removed")
        output_name = os.path.splitext(output_name)[0] + ".pdf"
        
        buffer = io.BytesIO()
        writer.write(buffer)
        buffer.seek(0)
        writer.close()
        
        response = send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=output_name
        )
        
        @after_this_request
        def cleanup(response):
            safe_remove(pdf_path)
            return response
        
        print(f"Deleted pages in {time.time() - start_time:.2f}s")
        return response
        
    except Exception as e:
        safe_remove(pdf_path)
        return jsonify({"error": f"Page removal failed: {str(e)}"}), 500

@app.route('/api/rotate-pdf', methods=['POST'])
def api_rotate_pdf():
    """Ultra-fast PDF rotation - FULLY IMPLEMENTED"""
    start_time = time.time()
    cleanup_temp()
    
    rotation = int(request.form.get('rotation', '90'))
    files = request.files.getlist('files')
    if not files or len(files) != 1:
        return jsonify({"error": "Upload exactly one PDF."}), 400
    
    paths = save_uploads(files)
    pdf_path = paths[0]
    
    if ext_of(pdf_path) not in ALLOWED_PDF_EXT:
        safe_remove(pdf_path)
        return jsonify({"error": "Only PDF files are allowed."}), 400

    try:
        # Read and rotate in one pass
        with open(pdf_path, 'rb') as f:
            reader = PdfReader(f)
        
        writer = PdfWriter()
        for page in reader.pages:
            page.rotate(rotation)
            writer.add_page(page)
        
        # Generate output
        original_name = secure_filename(files[0].filename)
        output_name = generate_unique_filename(original_name, f"rotated_{rotation}")
        output_name = os.path.splitext(output_name)[0] + ".pdf"
        
        buffer = io.BytesIO()
        writer.write(buffer)
        buffer.seek(0)
        writer.close()
        
        response = send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=output_name
        )
        
        @after_this_request
        def cleanup(response):
            safe_remove(pdf_path)
            return response
        
        print(f"Rotated PDF in {time.time() - start_time:.2f}s")
        return response
        
    except Exception as e:
        safe_remove(pdf_path)
        return jsonify({"error": f"Rotation failed: {str(e)}"}), 500

@app.route('/api/lock-pdf', methods=['POST'])
def api_lock_pdf():
    """Ultra-fast PDF encryption - FULLY IMPLEMENTED"""
    start_time = time.time()
    cleanup_temp()
    
    pin = request.form.get('pin', '').strip()
    if not pin or len(pin) != 4 or not pin.isdigit():
        return jsonify({"error": "PIN must be exactly 4 digits."}), 400
    
    files = request.files.getlist('files')
    if not files or len(files) != 1:
        return jsonify({"error": "Upload exactly one PDF."}), 400
    
    paths = save_uploads(files)
    pdf_path = paths[0]
    
    if ext_of(pdf_path) not in ALLOWED_PDF_EXT:
        safe_remove(pdf_path)
        return jsonify({"error": "Only PDF files are allowed."}), 400

    try:
        # Read and encrypt in one pass
        with open(pdf_path, 'rb') as f:
            reader = PdfReader(f)
        
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        
        # Encrypt with fast settings
        writer.encrypt(
            user_password=pin,
            owner_password=None,
            use_128bit=True,
            permissions_flag=0
        )
        
        # Generate output
        original_name = secure_filename(files[0].filename)
        output_name = generate_unique_filename(original_name, "locked")
        output_name = os.path.splitext(output_name)[0] + ".pdf"
        
        buffer = io.BytesIO()
        writer.write(buffer)
        buffer.seek(0)
        writer.close()
        
        response = send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=output_name
        )
        
        @after_this_request
        def cleanup(response):
            safe_remove(pdf_path)
            return response
        
        print(f"Locked PDF in {time.time() - start_time:.2f}s")
        return response
        
    except Exception as e:
        safe_remove(pdf_path)
        return jsonify({"error": f"Locking failed: {str(e)}"}), 500

@app.route('/api/unlock-pdf', methods=['POST'])
def api_unlock_pdf():
    """Ultra-fast PDF decryption - FULLY IMPLEMENTED"""
    start_time = time.time()
    cleanup_temp()
    
    password = request.form.get('password', '').strip()
    files = request.files.getlist('files')
    if not files or len(files) != 1:
        return jsonify({"error": "Upload exactly one PDF."}), 400
    if not password:
        return jsonify({"error": "Password is required."}), 400
    
    paths = save_uploads(files)
    pdf_path = paths[0]
    
    if ext_of(pdf_path) not in ALLOWED_PDF_EXT:
        safe_remove(pdf_path)
        return jsonify({"error": "Only PDF files are allowed."}), 400

    try:
        # Read and decrypt
        with open(pdf_path, 'rb') as f:
            reader = PdfReader(f)
        
        if reader.is_encrypted:
            if not reader.decrypt(password):
                safe_remove(pdf_path)
                return jsonify({"error": "Incorrect password."}), 400
        
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        
        # Generate output
        original_name = secure_filename(files[0].filename)
        output_name = generate_unique_filename(original_name, "unlocked")
        output_name = os.path.splitext(output_name)[0] + ".pdf"
        
        buffer = io.BytesIO()
        writer.write(buffer)
        buffer.seek(0)
        writer.close()
        
        response = send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=output_name
        )
        
        @after_this_request
        def cleanup(response):
            safe_remove(pdf_path)
            return response
        
        print(f"Unlocked PDF in {time.time() - start_time:.2f}s")
        return response
        
    except Exception as e:
        safe_remove(pdf_path)
        return jsonify({"error": f"Unlocking failed: {str(e)}"}), 500

# -----------------------------------------------------------------------------
# Tool APIs - Word Operations (ULTRA-FAST) - ALL IMPLEMENTED
# -----------------------------------------------------------------------------

@app.route('/api/word-to-pdf', methods=['POST'])
def api_word_to_pdf():
    """Ultra-fast Word to PDF conversion - FULLY IMPLEMENTED"""
    start_time = time.time()
    cleanup_temp()
    
    files = request.files.getlist('files')
    if not files or len(files) != 1:
        return jsonify({"error": "Upload exactly one Word file."}), 400

    paths = save_uploads(files)
    doc_path = paths[0]

    if ext_of(doc_path) not in ALLOWED_WORD_EXT:
        safe_remove(doc_path)
        return jsonify({"error": "Only DOC/DOCX files are supported."}), 400

    try:
        # Fast text extraction
        doc = Document(doc_path)
        text_content = []
        for para in doc.paragraphs:
            if para.text.strip():
                cleaned = clean_text_for_xml(para.text)
                if cleaned.strip():
                    text_content.append(cleaned.strip())
        
        text = "\n".join(text_content[:500])  # Limit text
        
        # Create PDF
        original_name = secure_filename(files[0].filename)
        output_name = generate_unique_filename(original_name, "converted_to_pdf")
        output_name = os.path.splitext(output_name)[0] + ".pdf"
        
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        width, height = letter
        left_margin = 50
        top = height - 50
        line_height = 14
        
        if text:
            # Fast text rendering
            paragraphs = text.split('\n\n')
            for para in paragraphs[:100]:  # Limit
                if para.strip():
                    lines = wrap_text(para, max_chars=95)
                    for line in lines:
                        c.drawString(left_margin, top, line)
                        top -= line_height
                        if top < 50:
                            c.showPage()
                            top = height - 50
                    top -= line_height / 2
                    if top < 50:
                        c.showPage()
                        top = height - 50
        
        c.save()
        buffer.seek(0)
        
        response = send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=output_name
        )
        
        @after_this_request
        def cleanup(response):
            safe_remove(doc_path)
            return response
        
        print(f"Word to PDF in {time.time() - start_time:.2f}s")
        return response
        
    except Exception as e:
        safe_remove(doc_path)
        return jsonify({"error": f"Conversion failed: {str(e)}"}), 500

@app.route('/api/merge-word', methods=['POST'])
def api_merge_word():
    """Ultra-fast Word document merging - FULLY IMPLEMENTED"""
    start_time = time.time()
    cleanup_temp()
    
    files = request.files.getlist('files')
    if not files or len(files) < 2:
        return jsonify({"error": "Upload at least two Word files."}), 400
    
    paths = save_uploads(files)
    for p in paths:
        if ext_of(p) not in ALLOWED_WORD_EXT:
            safe_remove_all(paths)
            return jsonify({"error": "Only DOC/DOCX files are allowed."}), 400

    try:
        # Merge documents efficiently
        merged = Document()
        
        for idx, doc_path in enumerate(paths):
            d = Document(doc_path)
            
            # Extract text efficiently
            paragraphs = []
            for para in d.paragraphs:
                if para.text.strip():
                    cleaned = clean_text_for_xml(para.text)
                    if cleaned.strip():
                        paragraphs.append(cleaned.strip())
            
            # Add to merged document with limits
            for para in paragraphs[:100]:  # Limit per document
                safe_add_paragraph(merged, para)
            
            # Add separator between documents
            if idx < len(paths) - 1:
                merged.add_paragraph("\n" + "=" * 50 + "\n")
        
        # Generate output
        original_name = secure_filename(files[0].filename)
        output_name = generate_unique_filename(original_name, "merged")
        output_name = os.path.splitext(output_name)[0] + ".docx"
        
        buffer = io.BytesIO()
        merged.save(buffer)
        buffer.seek(0)
        
        response = send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=output_name
        )
        
        @after_this_request
        def cleanup(response):
            safe_remove_all(paths)
            return response
        
        print(f"Merged {len(files)} Word docs in {time.time() - start_time:.2f}s")
        return response
        
    except Exception as e:
        safe_remove_all(paths)
        return jsonify({"error": f"Merging failed: {str(e)}"}), 500

@app.route('/api/word-to-text', methods=['POST'])
def api_word_to_text():
    """Ultra-fast Word to Text conversion - FULLY IMPLEMENTED"""
    start_time = time.time()
    cleanup_temp()
    
    files = request.files.getlist('files')
    if not files or len(files) != 1:
        return jsonify({"error": "Upload exactly one Word file."}), 400
    
    paths = save_uploads(files)
    doc_path = paths[0]
    
    if ext_of(doc_path) not in ALLOWED_WORD_EXT:
        safe_remove(doc_path)
        return jsonify({"error": "Only DOC/DOCX files are allowed."}), 400

    try:
        # Fast text extraction
        doc = Document(doc_path)
        text_content = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                cleaned = clean_text_for_xml(para.text)
                if cleaned.strip():
                    text_content.append(cleaned)
        
        # Generate output
        original_name = secure_filename(files[0].filename)
        output_name = generate_unique_filename(original_name, "extracted_text")
        output_name = os.path.splitext(output_name)[0] + ".txt"
        
        buffer = io.BytesIO('\n'.join(text_content).encode('utf-8'))
        buffer.seek(0)
        
        response = send_file(
            buffer,
            mimetype='text/plain',
            as_attachment=True,
            download_name=output_name
        )
        
        @after_this_request
        def cleanup(response):
            safe_remove(doc_path)
            return response
        
        print(f"Word to Text in {time.time() - start_time:.2f}s")
        return response
        
    except Exception as e:
        safe_remove(doc_path)
        return jsonify({"error": f"Conversion failed: {str(e)}"}), 500

# -----------------------------------------------------------------------------
# Tool APIs - Text Operations (ULTRA-FAST) - ALL IMPLEMENTED
# -----------------------------------------------------------------------------

@app.route('/api/text-to-pdf', methods=['POST'])
def api_text_to_pdf():
    """Ultra-fast Text to PDF conversion - FULLY IMPLEMENTED"""
    start_time = time.time()
    
    text = (request.form.get('text') or '').strip()
    if not text:
        return jsonify({"error": "Text content is required."}), 400

    cleaned_text = clean_text_for_xml(text)
    
    unique_id = str(uuid.uuid4())[:12]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_name = f"text_converted_{timestamp}_{unique_id}.pdf"

    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    left_margin = 50
    top = height - 50
    line_height = 14
    
    # Fast text processing
    lines = cleaned_text.splitlines()[:500]  # Limit lines
    for line in lines:
        if line.strip():
            for chunk in wrap_text(line, max_chars=95):
                c.drawString(left_margin, top, chunk)
                top -= line_height
                if top < 50:
                    c.showPage()
                    top = height - 50
        else:
            top -= line_height
            if top < 50:
                c.showPage()
                top = height - 50
    
    c.save()
    buffer.seek(0)
    
    print(f"Text to PDF in {time.time() - start_time:.2f}s")
    
    return send_file(
        buffer,
        mimetype='application/pdf',
        as_attachment=True,
        download_name=output_name
    )

@app.route('/api/text-to-word', methods=['POST'])
def api_text_to_word():
    """Ultra-fast Text to Word conversion - FULLY IMPLEMENTED"""
    start_time = time.time()
    
    text = (request.form.get('text') or '').strip()
    if not text:
        return jsonify({"error": "Text content is required."}), 400
    
    cleaned_text = clean_text_for_xml(text)
    
    unique_id = str(uuid.uuid4())[:12]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_name = f"text_converted_{timestamp}_{unique_id}.docx"
    
    doc = Document()
    
    if cleaned_text:
        lines = cleaned_text.splitlines()[:300]  # Limit lines
        for line in lines:
            if line.strip():
                safe_add_paragraph(doc, line)
    else:
        doc.add_paragraph("No text content provided.")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    print(f"Text to Word in {time.time() - start_time:.2f}s")
    
    return send_file(
        buffer,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=output_name
    )

# -----------------------------------------------------------------------------
# Tool APIs - Images to PDF (ULTRA-FAST) - ALL IMPLEMENTED
# -----------------------------------------------------------------------------

@app.route('/api/images-to-pdf', methods=['POST'])
def api_images_to_pdf():
    """Ultra-fast Images to PDF conversion - FULLY IMPLEMENTED"""
    start_time = time.time()
    cleanup_temp()
    
    files = request.files.getlist('files')
    if not files or len(files) < 1:
        return jsonify({"error": "Upload at least one image."}), 400
    
    paths = save_uploads(files)
    
    # Validate all are images
    for p in paths:
        if ext_of(p) not in ALLOWED_IMAGE_EXT:
            safe_remove_all(paths)
            return jsonify({"error": "Only image files (JPG, PNG, WEBP, BMP, TIFF, GIF) are allowed."}), 400

    try:
        # Process images in parallel
        def process_image(image_path):
            try:
                img = Image.open(image_path)
                img = ImageOps.exif_transpose(img)
                # Convert to RGB if necessary
                if img.mode in ('RGBA', 'LA', 'P'):
                    img = img.convert('RGB')
                return img
            except Exception as e:
                print(f"Failed to process image {image_path}: {e}")
                return None
        
        with ThreadPoolExecutor(max_workers=4) as executor:
            images = list(executor.map(process_image, paths))
        
        # Filter out None values
        images = [img for img in images if img is not None]
        
        if not images:
            safe_remove_all(paths)
            return jsonify({"error": "Failed to process images."}), 400
        
        # Create PDF
        original_name = secure_filename(files[0].filename)
        output_name = generate_unique_filename(original_name, "images_to_pdf")
        output_name = os.path.splitext(output_name)[0] + ".pdf"
        
        buffer = io.BytesIO()
        if len(images) == 1:
            images[0].save(buffer, format='PDF', save_all=True)
        else:
            first, rest = images[0], images[1:]
            first.save(buffer, format='PDF', save_all=True, append_images=rest)
        
        buffer.seek(0)
        
        response = send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=output_name
        )
        
        @after_this_request
        def cleanup(response):
            safe_remove_all(paths)
            return response
        
        print(f"Converted {len(images)} images to PDF in {time.time() - start_time:.2f}s")
        return response
        
    except Exception as e:
        safe_remove_all(paths)
        return jsonify({"error": f"Conversion failed: {str(e)}"}), 500

# -----------------------------------------------------------------------------
# Installation instructions endpoint
# -----------------------------------------------------------------------------

@app.route('/api/ocr-install', methods=['GET'])
def api_ocr_install():
    """Get OCR installation instructions"""
    instructions = {
        "ocr_available": OCR_AVAILABLE,
        "instructions": {
            "linux_ubuntu": [
                "sudo apt-get update",
                "sudo apt-get install -y tesseract-ocr",
                "sudo apt-get install -y libtesseract-dev",
                "sudo apt-get install -y poppler-utils",
                "pip install pytesseract pdf2image pillow"
            ],
            "macos": [
                "brew install tesseract",
                "brew install poppler",
                "pip install pytesseract pdf2image pillow"
            ],
            "windows": [
                "Download Tesseract OCR: https://github.com/UB-Mannheim/tesseract/wiki",
                "Add Tesseract to PATH: C:\\Program Files\\Tesseract-OCR",
                "pip install pytesseract pdf2image pillow"
            ],
            "python_packages": [
                "pip install pytesseract",
                "pip install pdf2image",
                "pip install pillow"
            ]
        }
    }
    return jsonify(instructions)

# -----------------------------------------------------------------------------
# Health check endpoint
# -----------------------------------------------------------------------------
@app.route('/health', methods=['GET'])
def health_check():
    return jsonify({
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "ocr_available": OCR_ENABLED,
        "cache_size": len(conversion_cache)
    }), 200

# -----------------------------------------------------------------------------
# Error handlers
# -----------------------------------------------------------------------------
@app.errorhandler(404)
def page_not_found(e):
    return jsonify({"error": "Page not found. Please check the URL."}), 404

@app.errorhandler(413)
def too_large(e):
    return jsonify({"error": f"File too large (max {MAX_CONTENT_LENGTH // (1024*1024)} MB)."}), 413

@app.errorhandler(400)
def bad_request(e):
    return jsonify({"error": str(e.description) if e.description else "Bad request."}), 400

@app.errorhandler(500)
def server_error(e):
    return jsonify({"error": "Internal server error."}), 500

# -----------------------------------------------------------------------------
# Static file serving
# -----------------------------------------------------------------------------
@app.route('/static/<path:filename>')
def serve_static(filename):
    return send_from_directory('static', filename)

# -----------------------------------------------------------------------------
# CORS Configuration
# -----------------------------------------------------------------------------
@app.after_request
def after_request(response):
    """Add CORS headers"""
    response.headers.add('Access-Control-Allow-Origin', '*')
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization,X-Requested-With')
    response.headers.add('Access-Control-Allow-Methods', 'GET,PUT,POST,DELETE,OPTIONS')
    response.headers.add('Access-Control-Allow-Credentials', 'true')
    return response

# -----------------------------------------------------------------------------
# Cache management
# -----------------------------------------------------------------------------
def cleanup_cache():
    """Periodically clean old cache entries"""
    while True:
        time.sleep(300)  # Run every 5 minutes
        try:
            current_time = time.time()
            expired_keys = []
            
            for key, (cache_time, _) in list(conversion_cache.items()):
                if current_time - cache_time > CACHE_TTL_SECONDS:
                    expired_keys.append(key)
            
            for key in expired_keys:
                del conversion_cache[key]
            
            # Keep cache size manageable
            if len(conversion_cache) > 200:
                keys = list(conversion_cache.keys())[:-200]
                for key in keys:
                    del conversion_cache[key]
        except:
            pass

# Start cleanup thread
cache_cleaner = threading.Thread(target=cleanup_cache, daemon=True)
cache_cleaner.start()

# -----------------------------------------------------------------------------
# Check OCR dependencies on startup
# -----------------------------------------------------------------------------
def check_ocr_dependencies():
    """Check and report OCR dependencies"""
    if not OCR_AVAILABLE:
        print("\n" + "="*60)
        print("OCR FEATURE SETUP REQUIRED")
        print("="*60)
        print("\nFor OCR functionality (handling scanned PDFs), install:")
        print("\n1. System dependencies (Ubuntu/Debian):")
        print("   sudo apt-get update")
        print("   sudo apt-get install -y tesseract-ocr")
        print("   sudo apt-get install -y libtesseract-dev")
        print("   sudo apt-get install -y poppler-utils")
        print("\n2. Python packages:")
        print("   pip install pytesseract")
        print("   pip install pdf2image")
        print("   pip install pillow")
        print("\n3. Language packs (optional):")
        print("   sudo apt-get install -y tesseract-ocr-eng  # English")
        print("   sudo apt-get install -y tesseract-ocr-spa  # Spanish")
        print("   sudo apt-get install -y tesseract-ocr-fra  # French")
        print("   sudo apt-get install -y tesseract-ocr-deu  # German")
        print("="*60 + "\n")
    else:
        print("\n✓ OCR functionality is AVAILABLE")
        print("  Scanned PDFs can be processed\n")
    
    return OCR_AVAILABLE

# -----------------------------------------------------------------------------
# Run the application
# -----------------------------------------------------------------------------
if __name__ == '__main__':
    print(f"\n🚀 Starting iMasterPDF with OCR Support")
    print(f"   OCR enabled: {OCR_ENABLED}")
    print(f"   Max workers: {MAX_WORKERS}")
    print(f"   Upload directory: {UPLOAD_DIR}")
    print(f"   Max file size: {MAX_CONTENT_LENGTH // (1024*1024)} MB")
    
    # Check dependencies
    check_ocr_dependencies()
    
    app.run(host='0.0.0.0', port=8000, debug=False, threaded=True)