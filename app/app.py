import os
import io
import shutil
import tempfile
import uuid
import re
import time
import hashlib
import zipfile
from datetime import datetime, timedelta
from concurrent.futures import ThreadPoolExecutor
from functools import lru_cache

from flask import Flask, render_template, send_file, request, abort, Response, jsonify, send_from_directory, after_this_request
from werkzeug.utils import secure_filename

from PyPDF2 import PdfReader, PdfWriter, PdfMerger
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4, letter
from PIL import Image
from pdfminer.high_level import extract_text
# -----------------------------------------------------------------------------
# OCR IMPORTS - ADD THESE HERE
# -----------------------------------------------------------------------------
# CORRECTED IMPORT (remove the problematic function)
from app.ocr import (
    pdf_to_word_with_ocr,
    pdf_to_text_with_ocr,
    image_to_text,
    image_to_word,
    extract_text_from_file,
    is_scanned_pdf
    # is_image_based_document  # Commented out for now
)

# -----------------------------------------------------------------------------
# Flask app configuration with correct template paths
# ----------------------------------------------------------------------------
# Get the base directory (where Dockerfile is)
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# Create Flask app with correct template and static paths
app = Flask(__name__, 
           template_folder=os.path.join(BASE_DIR, 'templates'),
           static_folder=os.path.join(BASE_DIR, 'static'))

print(f"✅ App initialized with:")
print(f"   Base directory: {BASE_DIR}")
print(f"   Templates: {os.path.join(BASE_DIR, 'templates')}")
print(f"   Static files: {os.path.join(BASE_DIR, 'static')}")

# -----------------------------------------------------------------------------
# Performance settings
# -----------------------------------------------------------------------------
MAX_CONTENT_LENGTH = 50 * 1024 * 1024  # 50 MB per file
UPLOAD_DIR = os.path.join(tempfile.gettempdir(), "imasterpdf_uploads")
OUTPUT_DIR = os.path.join(tempfile.gettempdir(), "imasterpdf_outputs")
CLEANUP_AGE_MINUTES = 30
MAX_WORKERS = 4  # For parallel processing
MAX_PAGES_TO_EXTRACT = 100  # Limit for large PDFs
CACHE_ENABLED = True

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH

# Thread pool for parallel processing
executor = ThreadPoolExecutor(max_workers=MAX_WORKERS)

# Cache for repeated conversions
conversion_cache = {}

ALLOWED_IMAGE_EXT = {'.jpg', '.jpeg', '.png', '.webp', '.bmp', '.tiff', '.tif'}
ALLOWED_PDF_EXT = {'.pdf'} 
ALLOWED_WORD_EXT = {'.docx', '.doc'}
ALLOWED_TEXT_EXT = {'.txt'}

# -----------------------------------------------------------------------------
# Performance optimization utilities
# -----------------------------------------------------------------------------
def clean_text_for_xml(text):
    """
    Clean text to make it XML compatible.
    Removes NULL bytes, control characters, and other problematic chars.
    """
    if not text:
        return ""
    
    # Remove NULL bytes
    text = text.replace('\x00', '')
    
    # Remove other control characters (except common whitespace: \n, \t, \r)
    text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]', '', text)
    
    # Replace other problematic Unicode characters
    replacements = {
        '\u2028': ' ',
        '\u2029': ' ',
        '\uFEFF': '',
    }
    
    for old, new in replacements.items():
        text = text.replace(old, new)
    
    # Ensure valid UTF-8
    try:
        text = text.encode('utf-8', 'ignore').decode('utf-8')
    except:
        text = text.encode('ascii', 'ignore').decode('ascii')
    
    return text

def safe_add_paragraph(doc, text):
    """
    Safely add a paragraph to a Word document, handling any XML errors.
    """
    try:
        cleaned_text = clean_text_for_xml(text)
        if cleaned_text.strip():
            doc.add_paragraph(cleaned_text.strip())
    except Exception:
        pass

def get_file_hash(file_path):
    """Get MD5 hash of file for caching"""
    hasher = hashlib.md5()
    with open(file_path, 'rb') as f:
        buf = f.read(65536)
        while len(buf) > 0:
            hasher.update(buf)
            buf = f.read(65536)
    return hasher.hexdigest()

def fast_extract_text(pdf_path):
    """
    Fast text extraction with multiple optimized methods.
    """
    start_time = time.time()
    file_size = os.path.getsize(pdf_path)
    
    # For very small files, use simple extraction
    if file_size < 102400:  # < 100KB
        try:
            with open(pdf_path, 'rb') as file:
                reader = PdfReader(file)
                text = ""
                for page in reader.pages[:MAX_PAGES_TO_EXTRACT]:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                if text.strip():
                    print(f"Fast PyPDF2 extraction: {time.time() - start_time:.2f}s")
                    return clean_text_for_xml(text)
        except:
            pass
    
    # For larger files, try parallel extraction
    try:
        text = parallel_pdf_extraction(pdf_path)
        if text and len(text.strip()) > 50:
            print(f"Parallel extraction: {time.time() - start_time:.2f}s")
            return clean_text_for_xml(text)
    except:
        pass
    
    # Fallback to optimized pdfminer
    try:
        # Extract only first N pages for speed
        text = extract_text(
            pdf_path,
            maxpages=MAX_PAGES_TO_EXTRACT,
            caching=True,
            laparams=None  # Disable layout analysis for speed
        )
        print(f"Optimized pdfminer extraction: {time.time() - start_time:.2f}s")
        return clean_text_for_xml(text or "")
    except Exception as e:
        print(f"All extraction methods failed: {e}")
        return ""

def parallel_pdf_extraction(pdf_path, max_workers=MAX_WORKERS):
    """
    Extract text from PDF pages in parallel for speed.
    """
    try:
        with open(pdf_path, 'rb') as file:
            reader = PdfReader(file)
            pages = reader.pages[:MAX_PAGES_TO_EXTRACT]  # Limit pages
            
            def extract_page(page):
                try:
                    return page.extract_text() or ""
                except:
                    return ""
            
            # Extract pages in parallel
            texts = list(executor.map(extract_page, pages))
            return "\n".join(texts)
    except Exception as e:
        print(f"Parallel extraction failed: {e}")
        return ""

def optimize_pdf_for_extraction(pdf_path):
    """
    Optimize PDF for faster text extraction.
    Returns optimized file path or original if optimization fails.
    """
    try:
        file_size = os.path.getsize(pdf_path)
        
        # Skip optimization for small files
        if file_size < 5 * 1024 * 1024:  # < 5MB
            return pdf_path
            
        # For large files, extract only first N pages
        with open(pdf_path, 'rb') as file:
            reader = PdfReader(file)
            if len(reader.pages) <= MAX_PAGES_TO_EXTRACT:
                return pdf_path
            
            # Create optimized PDF with only first N pages
            writer = PdfWriter()
            for i in range(min(MAX_PAGES_TO_EXTRACT, len(reader.pages))):
                writer.add_page(reader.pages[i])
            
            optimized_path = pdf_path + "_optimized.pdf"
            with open(optimized_path, 'wb') as f:
                writer.write(f)
            
            return optimized_path
    except:
        return pdf_path

# -----------------------------------------------------------------------------
# Utility helpers
# -----------------------------------------------------------------------------
def ext_of(filename):
    return os.path.splitext(filename.lower())[1]

def validate_file(stream):
    stream.seek(0, os.SEEK_END)
    size = stream.tell()
    stream.seek(0)
    if size < 1024:
        abort(Response("File too small (min 1 KB).", status=400))
    if size > MAX_CONTENT_LENGTH:
        abort(Response("File too large (max 50 MB).", status=400))

def generate_unique_filename(original_filename, suffix=""):
    """Generate a unique filename with UUID and timestamp"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    unique_id = str(uuid.uuid4())[:12]
    name, ext = os.path.splitext(original_filename)
    safe_name = secure_filename(name)
    
    if suffix:
        return f"{safe_name}_{suffix}_{timestamp}_{unique_id}{ext}"
    return f"{safe_name}_{timestamp}_{unique_id}{ext}"

def save_uploads(files):
    saved = []
    for storage in files:
        validate_file(storage.stream)
        filename = secure_filename(storage.filename)
        if not filename:
            abort(Response("Invalid filename.", status=400))
        
        unique_filename = generate_unique_filename(storage.filename)
        path = os.path.join(UPLOAD_DIR, unique_filename)
        storage.save(path)
        saved.append(path)
    return saved

def cleanup_temp():
    cutoff = datetime.utcnow() - timedelta(minutes=CLEANUP_AGE_MINUTES)
    for base in (UPLOAD_DIR, OUTPUT_DIR):
        for name in os.listdir(base):
            path = os.path.join(base, name)
            try:
                mtime = datetime.utcfromtimestamp(os.path.getmtime(path))
                if mtime < cutoff:
                    if os.path.isdir(path):
                        shutil.rmtree(path, ignore_errors=True)
                    else:
                        os.remove(path)
            except Exception:
                pass

def wrap_text(text, max_chars=95):
    words = text.split(' ')
    lines, current = [], []
    length = 0
    for w in words:
        add_len = len(w) + (1 if current else 0)
        if length + add_len <= max_chars:
            current.append(w)
            length += add_len
        else:
            lines.append(' '.join(current))
            current = [w]
            length = len(w)
    if current:
        lines.append(' '.join(current))
    return lines

def parse_pages(pages_str):
    pages = set()
    parts = [p.strip() for p in pages_str.split(',') if p.strip()]
    for part in parts:
        if '-' in part:
            a, b = part.split('-', 1)
            try:
                start = int(a); end = int(b)
                for i in range(min(start, end), max(start, end)+1):
                    pages.add(i)
            except ValueError:
                abort(Response("Invalid page range.", status=400))
        else:
            try:
                pages.add(int(part))
            except ValueError:
                abort(Response("Invalid page number.", status=400))
    return pages

def safe_remove(path):
    try:
        if os.path.exists(path):
            os.remove(path)
    except Exception:
        pass

def safe_remove_all(paths):
    for path in paths:
        safe_remove(path)

# -----------------------------------------------------------------------------
# SPA Routes for each tool page
# -----------------------------------------------------------------------------

@app.route('/')
@app.route('/index')
@app.route('/index.html')
def index():
    """Main landing page"""
    return render_template('index.html')

@app.route('/split')
@app.route('/split.html')
def split_pdf():
    """Split PDF tool page"""
    return render_template('split.html')

@app.route('/mergepdf')
@app.route('/mergepdf.html')
def merge_pdf():
    """Merge PDF tool page"""
    return render_template('mergepdf.html')

@app.route('/deletepdf')
@app.route('/deletepdf.html')
def delete_pdf():
    """Delete pages from PDF tool page"""
    return render_template('deletepdf.html')

@app.route('/rotatepdf')
@app.route('/rotatepdf.html')
def rotate_pdf():
    """Rotate PDF pages tool page"""
    return render_template('rotatepdf.html')

@app.route('/pdftoword')
@app.route('/pdftoword.html')
def pdf_to_word():
    """PDF to Word converter page"""
    return render_template('pdftoword.html')

@app.route('/lockpdf')
@app.route('/lockpdf.html')
def lock_pdf():
    """Lock PDF with password page"""
    return render_template('lockpdf.html')

@app.route('/unlockpdf')
@app.route('/unlockpdf.html')
def unlock_pdf():
    """Unlock PDF page"""
    return render_template('unlockpdf.html')

@app.route('/wordtopdf')
@app.route('/wordtopdf.html')
def word_to_pdf():
    """Word to PDF converter page"""
    return render_template('wordtopdf.html')

@app.route('/mergeword')
@app.route('/mergeword.html')
def merge_word():
    """Merge Word documents page"""
    return render_template('mergeword.html')

@app.route('/wordtotext')
@app.route('/wordtotext.html')
def word_to_text():
    """Word to Text converter page"""
    return render_template('wordtotext.html')

@app.route('/texttopdf')
@app.route('/texttopdf.html')
def text_to_pdf():
    """Text to PDF converter page"""
    return render_template('texttopdf.html')

@app.route('/texttoword')
@app.route('/texttoword.html')
def text_to_word():
    """Text to Word converter page"""
    return render_template('texttoword.html')

@app.route('/imagestopdf')
@app.route('/imagestopdf.html')
def images_to_pdf():
    """Images to PDF converter page"""
    return render_template('imagestopdf.html')

# -----------------------------------------------------------------------------
# Catch-all route for other .html files
# -----------------------------------------------------------------------------
@app.route('/<path:filename>.html')
def serve_html(filename):
    """Catch-all route for any .html file requests"""
    try:
        return render_template(f'{filename}.html')
    except:
        abort(404)

# -----------------------------------------------------------------------------
# Contact API
# -----------------------------------------------------------------------------
@app.route('/api/contact', methods=['POST'])
def api_contact():
    data = request.get_json(silent=True) or {}
    name = (data.get('name') or '').strip()
    email = (data.get('email') or '').strip()
    message = (data.get('message') or '').strip()
    if not name or not email or not message:
        return jsonify({"error": "Please provide name, email, and message."}), 400
    return jsonify({"status": "ok", "received": {"name": name, "email": email}}), 200

# -----------------------------------------------------------------------------
# Tool APIs - PDF Operations
# -----------------------------------------------------------------------------

@app.route('/api/merge-pdf', methods=['POST'])
def api_merge_pdf():
    cleanup_temp()
    files = request.files.getlist('files')
    if not files or len(files) < 2:
        return jsonify({"error": "Upload at least two PDFs."}), 400
    
    paths = save_uploads(files)
    for p in paths:
        if ext_of(p) not in ALLOWED_PDF_EXT:
            safe_remove_all(paths)
            return jsonify({"error": "Only PDF files are allowed."}), 400

    merger = PdfMerger()
    try:
        for p in paths:
            merger.append(p)
        
        original_name = secure_filename(files[0].filename)
        output_name = generate_unique_filename(original_name, "merged")
        output_name = os.path.splitext(output_name)[0] + ".pdf"
        
        buffer = io.BytesIO()
        merger.write(buffer)
        buffer.seek(0)
        
        response = send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=output_name
        )
        
        @after_this_request
        def cleanup(response):
            safe_remove_all(paths)
            merger.close()
            return response
        
        return response
        
    except Exception as e:
        safe_remove_all(paths)
        merger.close()
        return jsonify({"error": f"Merging failed: {str(e)}"}), 500

@app.route('/api/split-pdf', methods=['POST'])
def api_split_pdf():
    """Split PDF by page ranges"""
    cleanup_temp()
    files = request.files.getlist('files')
    if not files or len(files) != 1:
        return jsonify({"error": "Upload exactly one PDF."}), 400
    
    ranges_str = request.form.get('ranges', '').strip()
    if not ranges_str:
        return jsonify({"error": "Page ranges are required."}), 400
    
    paths = save_uploads(files)
    pdf_path = paths[0]
    if ext_of(pdf_path) not in ALLOWED_PDF_EXT:
        safe_remove(pdf_path)
        return jsonify({"error": "Only PDF files are allowed."}), 400
    
    try:
        reader = PdfReader(pdf_path)
        total_pages = len(reader.pages)
        
        ranges = []
        parts = [p.strip() for p in ranges_str.split(',') if p.strip()]
        for part in parts:
            if '-' in part:
                start, end = part.split('-', 1)
                try:
                    start = int(start); end = int(end)
                    if 1 <= start <= total_pages and 1 <= end <= total_pages:
                        ranges.append((min(start, end)-1, max(start, end)))
                    else:
                        safe_remove(pdf_path)
                        return jsonify({"error": f"Page range out of bounds (1-{total_pages})."}), 400
                except ValueError:
                    safe_remove(pdf_path)
                    return jsonify({"error": "Invalid page range format."}), 400
            else:
                try:
                    page = int(part)
                    if 1 <= page <= total_pages:
                        ranges.append((page-1, page))
                    else:
                        safe_remove(pdf_path)
                        return jsonify({"error": f"Page out of bounds (1-{total_pages})."}), 400
                except ValueError:
                    safe_remove(pdf_path)
                    return jsonify({"error": "Invalid page number."}), 400
        
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for i, (start_idx, end_page) in enumerate(ranges):
                writer = PdfWriter()
                for page_idx in range(start_idx, end_page):
                    writer.add_page(reader.pages[page_idx])
                
                split_buffer = io.BytesIO()
                writer.write(split_buffer)
                split_buffer.seek(0)
                
                original_name = secure_filename(files[0].filename)
                split_name = generate_unique_filename(original_name, f"split_{i+1}")
                split_name = os.path.splitext(split_name)[0] + ".pdf"
                
                zipf.writestr(split_name, split_buffer.getvalue())
                writer.close()
        
        zip_buffer.seek(0)
        original_name = secure_filename(files[0].filename)
        zip_name = generate_unique_filename(original_name, "split_parts")
        zip_name = os.path.splitext(zip_name)[0] + ".zip"
        
        response = send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name=zip_name
        )
        
        @after_this_request
        def cleanup(response):
            safe_remove(pdf_path)
            return response
        
        return response
        
    except Exception as e:
        safe_remove(pdf_path)
        return jsonify({"error": f"Splitting failed: {str(e)}"}), 500

@app.route('/api/delete-pages-pdf', methods=['POST'])
def api_delete_pages_pdf():
    cleanup_temp()
    pages_str = request.form.get('pages', '').strip()
    files = request.files.getlist('files')
    if not files or len(files) != 1:
        return jsonify({"error": "Upload exactly one PDF."}), 400
    if not pages_str:
        return jsonify({"error": "Pages to remove are required."}), 400
    
    paths = save_uploads(files)
    pdf_path = paths[0]
    if ext_of(pdf_path) not in ALLOWED_PDF_EXT:
        safe_remove(pdf_path)
        return jsonify({"error": "Only PDF files are allowed."}), 400

    pages_to_remove = parse_pages(pages_str)
    writer = PdfWriter()
    
    try:
        reader = PdfReader(pdf_path)
        total = len(reader.pages)
        
        original_name = secure_filename(files[0].filename)
        output_name = generate_unique_filename(original_name, "pages_removed")
        output_name = os.path.splitext(output_name)[0] + ".pdf"
        
        for i in range(total):
            if (i+1) not in pages_to_remove:
                writer.add_page(reader.pages[i])
        
        buffer = io.BytesIO()
        writer.write(buffer)
        buffer.seek(0)
        
        response = send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=output_name
        )
        
        @after_this_request
        def cleanup(response):
            safe_remove(pdf_path)
            writer.close()
            return response
        
        return response
        
    except Exception as e:
        safe_remove(pdf_path)
        writer.close()
        return jsonify({"error": f"Page removal failed: {str(e)}"}), 500

@app.route('/api/rotate-pdf', methods=['POST'])
def api_rotate_pdf():
    cleanup_temp()
    rotation = int(request.form.get('rotation', '90'))
    files = request.files.getlist('files')
    if not files or len(files) != 1:
        return jsonify({"error": "Upload exactly one PDF."}), 400
    
    paths = save_uploads(files)
    pdf_path = paths[0]
    if ext_of(pdf_path) not in ALLOWED_PDF_EXT:
        safe_remove(pdf_path)
        return jsonify({"error": "Only PDF files are allowed."}), 400

    writer = PdfWriter()
    try:
        reader = PdfReader(pdf_path)
        
        original_name = secure_filename(files[0].filename)
        output_name = generate_unique_filename(original_name, f"rotated_{rotation}")
        output_name = os.path.splitext(output_name)[0] + ".pdf"
        
        for page in reader.pages:
            page.rotate(rotation)
            writer.add_page(page)
        
        buffer = io.BytesIO()
        writer.write(buffer)
        buffer.seek(0)
        
        response = send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=output_name
        )
        
        @after_this_request
        def cleanup(response):
            safe_remove(pdf_path)
            writer.close()
            return response
        
        return response
        
    except Exception as e:
        safe_remove(pdf_path)
        writer.close()
        return jsonify({"error": f"Rotation failed: {str(e)}"}), 500

@app.route('/api/lock-pdf', methods=['POST'])
def api_lock_pdf():
    cleanup_temp()
    pin = request.form.get('pin', '').strip()
    if not pin or len(pin) != 4 or not pin.isdigit():
        return jsonify({"error": "PIN must be exactly 4 digits."}), 400
    
    files = request.files.getlist('files')
    if not files or len(files) != 1:
        return jsonify({"error": "Upload exactly one PDF."}), 400
    
    paths = save_uploads(files)
    pdf_path = paths[0]
    if ext_of(pdf_path) not in ALLOWED_PDF_EXT:
        safe_remove(pdf_path)
        return jsonify({"error": "Only PDF files are allowed."}), 400

    writer = PdfWriter()
    try:
        reader = PdfReader(pdf_path)
        
        original_name = secure_filename(files[0].filename)
        output_name = generate_unique_filename(original_name, "locked")
        output_name = os.path.splitext(output_name)[0] + ".pdf"
        
        for page in reader.pages:
            writer.add_page(page)
        writer.encrypt(pin)
        
        buffer = io.BytesIO()
        writer.write(buffer)
        buffer.seek(0)
        
        response = send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=output_name
        )
        
        @after_this_request
        def cleanup(response):
            safe_remove(pdf_path)
            writer.close()
            return response
        
        return response
        
    except Exception as e:
        safe_remove(pdf_path)
        writer.close()
        return jsonify({"error": f"Locking failed: {str(e)}"}), 500

@app.route('/api/unlock-pdf', methods=['POST'])
def api_unlock_pdf():
    cleanup_temp()
    password = request.form.get('password', '').strip()
    files = request.files.getlist('files')
    if not files or len(files) != 1:
        return jsonify({"error": "Upload exactly one PDF."}), 400
    if not password:
        return jsonify({"error": "Password is required."}), 400
    
    paths = save_uploads(files)
    pdf_path = paths[0]
    if ext_of(pdf_path) not in ALLOWED_PDF_EXT:
        safe_remove(pdf_path)
        return jsonify({"error": "Only PDF files are allowed."}), 400

    writer = PdfWriter()
    try:
        reader = PdfReader(pdf_path)
        
        original_name = secure_filename(files[0].filename)
        output_name = generate_unique_filename(original_name, "unlocked")
        output_name = os.path.splitext(output_name)[0] + ".pdf"
        
        if reader.is_encrypted:
            if not reader.decrypt(password):
                safe_remove(pdf_path)
                writer.close()
                return jsonify({"error": "Incorrect password."}), 400
        for page in reader.pages:
            writer.add_page(page)
        
        buffer = io.BytesIO()
        writer.write(buffer)
        buffer.seek(0)
        
        response = send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=output_name
        )
        
        @after_this_request
        def cleanup(response):
            safe_remove(pdf_path)
            writer.close()
            return response
        
        return response
        
    except Exception as e:
        safe_remove(pdf_path)
        writer.close()
        return jsonify({"error": f"Unlocking failed: {str(e)}"}), 500

# -----------------------------------------------------------------------------
# Tool APIs - PDF to Word (OPTIMIZED for speed)
# -----------------------------------------------------------------------------


@app.route('/api/pdf-to-word', methods=['POST'])
def api_pdf_to_word():
    start_time = time.time()
    files = request.files.getlist('files')

    if not files or len(files) != 1:
        return jsonify({"error": "Upload exactly one PDF."}), 400

    pdf_file = files[0]
    pdf_path = os.path.join("uploads", pdf_file.filename)
    pdf_file.save(pdf_path)

    output_name = os.path.splitext(pdf_file.filename)[0] + ".docx"

    try:
        # =========================
        # CASE 1: SCANNED PDF → OCR
        # =========================
        if is_scanned_pdf(pdf_path):
            print("🔍 Scanned PDF detected → OCR")

            temp_docx = os.path.join(
                tempfile.gettempdir(),
                f"ocr_{uuid.uuid4().hex}.docx"
            )

            pdf_to_word_with_ocr(pdf_path, temp_docx)

            with open(temp_docx, "rb") as f:
                buffer = io.BytesIO(f.read())
                buffer.seek(0)

            os.remove(temp_docx)
            os.remove(pdf_path)

            return send_file(
                buffer,
                as_attachment=True,
                download_name=output_name,
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        # =========================
        # CASE 2: NORMAL TEXT PDF
        # =========================
        text = extract_text(pdf_path)

        doc = Document()
        doc.add_paragraph(text)   # IMPORTANT: add full text directly

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        os.remove(pdf_path)

        print(f"PDF → Word done in {time.time() - start_time:.2f}s")

        return send_file(
            buffer,
            as_attachment=True,
            download_name=output_name,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        if os.path.exists(pdf_path):
            os.remove(pdf_path)
        return jsonify({"error": str(e)}), 500
# -----------------------------------------------------------------------------
# Tool APIs - Word Operations (OPTIMIZED)
# -----------------------------------------------------------------------------

@app.route('/api/word-to-pdf', methods=['POST'])
def api_word_to_pdf():
    cleanup_temp()
    files = request.files.getlist('files')
    if not files or len(files) != 1:
        return jsonify({"error": "Upload exactly one Word file."}), 400

    paths = save_uploads(files)
    doc_path = paths[0]

    if ext_of(doc_path) not in ALLOWED_WORD_EXT:
        safe_remove(doc_path)
        return jsonify({"error": "Only DOC/DOCX files are supported."}), 400

    try:
        original_name = secure_filename(files[0].filename)
        output_name = generate_unique_filename(original_name, "converted_to_pdf")
        output_name = os.path.splitext(output_name)[0] + ".pdf"
        
        # Extract text (with OCR if needed)
        text = extract_text_from_file(doc_path)
        
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        width, height = letter
        left_margin = 50
        top = height - 50
        line_height = 14
        
        if text:
            paragraphs = text.split('\n\n')
            for para in paragraphs[:200]:
                if para.strip():
                    lines = wrap_text(para, max_chars=95)
                    for line in lines:
                        c.drawString(left_margin, top, line)
                        top -= line_height
                        if top < 50:
                            c.showPage()
                            top = height - 50
                    top -= line_height / 2
                    if top < 50:
                        c.showPage()
                        top = height - 50
        
        c.save()
        buffer.seek(0)
        
        response = send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=output_name
        )
        
        @after_this_request
        def cleanup(response):
            safe_remove(doc_path)
            return response
        
        return response
        
    except Exception as e:
        safe_remove(doc_path)
        return jsonify({"error": f"Conversion failed: {str(e)}"}), 500

@app.route('/api/word-to-text', methods=['POST'])
def api_word_to_text():
    cleanup_temp()
    files = request.files.getlist('files')
    if not files or len(files) != 1:
        return jsonify({"error": "Upload exactly one Word file."}), 400
    
    paths = save_uploads(files)
    doc_path = paths[0]
    if ext_of(doc_path) not in ALLOWED_WORD_EXT:
        safe_remove(doc_path)
        return jsonify({"error": "Only DOC/DOCX files are allowed."}), 400

    try:
        original_name = secure_filename(files[0].filename)
        output_name = generate_unique_filename(original_name, "extracted_text")
        output_name = os.path.splitext(output_name)[0] + ".txt"
        
        # OCR CHECK for image-based Word documents
        if is_image_based_document(doc_path):
            print("🖼️ Image-based Word detected, extracting text with OCR...")
            text = extract_text_from_file(doc_path)
        else:
            # Normal text extraction
            doc = Document(doc_path)
            text_content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    cleaned = clean_text_for_xml(para.text)
                    if cleaned.strip():
                        text_content.append(cleaned)
            text = '\n'.join(text_content)
        
        buffer = io.BytesIO(text.encode('utf-8'))
        buffer.seek(0)
        
        response = send_file(
            buffer,
            mimetype='text/plain',
            as_attachment=True,
            download_name=output_name
        )
        
        @after_this_request
        def cleanup(response):
            safe_remove(doc_path)
            return response
        
        return response
        
    except Exception as e:
        safe_remove(doc_path)
        return jsonify({"error": f"Conversion failed: {str(e)}"}), 500
# -----------------------------------------------------------------------------
# Tool APIs - Text Operations
# -----------------------------------------------------------------------------

@app.route('/api/text-to-pdf', methods=['POST'])
def api_text_to_pdf():
    cleanup_temp()
    text = (request.form.get('text') or '').strip()
    if not text:
        return jsonify({"error": "Text content is required."}), 400

    cleaned_text = clean_text_for_xml(text)
    
    unique_id = str(uuid.uuid4())[:12]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_name = f"text_converted_{timestamp}_{unique_id}.pdf"

    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    left_margin = 50
    top = height - 50
    line_height = 14
    
    lines = cleaned_text.splitlines()
    for line in lines[:500]:  # Limit for speed
        if line.strip():
            for chunk in wrap_text(line, max_chars=95):
                c.drawString(left_margin, top, chunk)
                top -= line_height
                if top < 50:
                    c.showPage()
                    top = height - 50
        else:
            top -= line_height
            if top < 50:
                c.showPage()
                top = height - 50
    
    c.save()
    buffer.seek(0)
    
    return send_file(
        buffer,
        mimetype='application/pdf',
        as_attachment=True,
        download_name=output_name
    )

@app.route('/api/text-to-word', methods=['POST'])
def api_text_to_word():
    cleanup_temp()
    text = (request.form.get('text') or '').strip()
    if not text:
        return jsonify({"error": "Text content is required."}), 400
    
    cleaned_text = clean_text_for_xml(text)
    
    unique_id = str(uuid.uuid4())[:12]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_name = f"text_converted_{timestamp}_{unique_id}.docx"
    
    doc = Document()
    
    if cleaned_text:
        lines = cleaned_text.splitlines()
        for line in lines[:500]:  # Limit for speed
            if line.strip():
                safe_add_paragraph(doc, line)
    else:
        doc.add_paragraph("No text content provided.")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return send_file(
        buffer,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=output_name
    )

# -----------------------------------------------------------------------------
# Tool APIs - Images to PDF
# -----------------------------------------------------------------------------

@app.route('/api/images-to-pdf', methods=['POST'])
def api_images_to_pdf():
    cleanup_temp()
    files = request.files.getlist('files')
    if not files or len(files) < 1:
        return jsonify({"error": "Upload at least one image."}), 400
    
    paths = save_uploads(files)
    for p in paths:
        if ext_of(p) not in ALLOWED_IMAGE_EXT:
            safe_remove_all(paths)
            return jsonify({"error": "Only image files (JPG, PNG, WEBP, BMP, TIFF) are allowed."}), 400

    try:
        original_name = secure_filename(files[0].filename)
        output_name = generate_unique_filename(original_name, "images_to_pdf")
        output_name = os.path.splitext(output_name)[0] + ".pdf"
        
        images = []
        for p in paths:
            img = Image.open(p).convert('RGB')
            images.append(img)
        
        buffer = io.BytesIO()
        if len(images) == 1:
            images[0].save(buffer, format='PDF', save_all=True)
        else:
            first, rest = images[0], images[1:]
            first.save(buffer, format='PDF', save_all=True, append_images=rest)
        
        buffer.seek(0)
        
        response = send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=output_name
        )
        
        @after_this_request
        def cleanup(response):
            safe_remove_all(paths)
            return response
        
        return response
        
    except Exception as e:
        safe_remove_all(paths)
        return jsonify({"error": f"Conversion failed: {str(e)}"}), 500

# -----------------------------------------------------------------------------
# Health check endpoint
# -----------------------------------------------------------------------------
@app.route('/health', methods=['GET'])
def health_check():
    return jsonify({"status": "healthy", "timestamp": datetime.now().isoformat()}), 200

# -----------------------------------------------------------------------------
# Error handlers
# -----------------------------------------------------------------------------
@app.errorhandler(404)
def page_not_found(e):
    return jsonify({"error": "Page not found. Please check the URL."}), 404

@app.errorhandler(413)
def too_large(e):
    return jsonify({"error": "File too large (max 50 MB)."}), 413

@app.errorhandler(400)
def bad_request(e):
    return jsonify({"error": str(e.description) if e.description else "Bad request."}), 400

@app.errorhandler(500)
def server_error(e):
    return jsonify({"error": "Internal server error."}), 500

# -----------------------------------------------------------------------------
# Static file serving for templates (if needed)
# -----------------------------------------------------------------------------
@app.route('/static/<path:filename>')
def serve_static(filename):
    return send_from_directory('static', filename)

# -----------------------------------------------------------------------------
# CORS Configuration for Render
# -----------------------------------------------------------------------------
@app.after_request
def after_request(response):
    """Add CORS headers for Render deployment"""
    response.headers.add('Access-Control-Allow-Origin', '*')
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
    response.headers.add('Access-Control-Allow-Methods', 'GET,PUT,POST,DELETE,OPTIONS')
    return response

# -----------------------------------------------------------------------------
# Cleanup thread for cache management
# -----------------------------------------------------------------------------
def cleanup_cache():
    """Periodically clean old cache entries"""
    while True:
        time.sleep(300)  # Run every 5 minutes
        try:
            # Keep only last 100 entries
            if len(conversion_cache) > 100:
                keys = list(conversion_cache.keys())[:-100]
                for key in keys:
                    del conversion_cache[key]
        except:
            pass

# Start cleanup thread in background
import threading
cache_cleaner = threading.Thread(target=cleanup_cache, daemon=True)
cache_cleaner.start()

# -----------------------------------------------------------------------------
# Run the application (for local development only)
# -----------------------------------------------------------------------------
if __name__ == '__main__':
    print(f"Starting iMasterPDF with {MAX_WORKERS} workers")
    print(f"Cache enabled: {CACHE_ENABLED}")
    print(f"Max pages to extract: {MAX_PAGES_TO_EXTRACT}")
    app.run(host='0.0.0.0', port=8000, debug=False, threaded=True)